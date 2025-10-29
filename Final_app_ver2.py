# Final_app_ver2.py
# Streamlit app with Workspace selection + full interactive loops (M2~M5) + M3 bulk-edit flow + M6
# pip install streamlit openai biopython python-docx beautifulsoup4 selenium pandas matplotlib scikit-learn pyarrow
# Run: streamlit run Final_app_ver2.py
# pyright: ignore[reportMissingImports]

import os, re, sys, json, time, importlib.util, subprocess, zipfile, gzip, shutil, threading, queue
from pathlib import Path
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="연구자동화 파이프라인", layout="wide")

# ----------------------- CODE ROOT -----------------------
try:
    ROOT = Path(__file__).resolve().parent
except NameError:
    ROOT = Path.cwd()  # very rare fallback

# ----------------------- Sidebar: Workspace -----------------------
st.sidebar.title("환경 설정")
st.sidebar.subheader("작업 폴더(Workspace)")

DEFAULT_HINT = (ROOT / "workspace").as_posix()
ws_input = st.sidebar.text_input(
    "경로 입력(없으면 생성)", value="", placeholder=DEFAULT_HINT  # 처음엔 공란
)

if st.sidebar.button("작업 폴더 적용"):
    # 입력이 비어 있으면 placeholder 경로로 생성
    wp = Path((ws_input or DEFAULT_HINT)).expanduser().resolve()
    wp.mkdir(parents=True, exist_ok=True)
    st.session_state.WORK_DIR = str(wp)
    st.sidebar.success(f"작업 폴더 설정됨: {wp}")

# 사용자가 아직 적용하지 않았다면 WORK_DIR 없음
WORK_DIR = (
    Path(st.session_state["WORK_DIR"]).resolve()
    if "WORK_DIR" in st.session_state
    else None
)

# 하위 출력 경로들은 WORK_DIR 확정 후에만 생성
RESULT_DIR = FIG_DIR = SUB_DIR = M4_OUT = None
if WORK_DIR:
    RESULT_DIR = WORK_DIR / "analysis_result"
    FIG_DIR = RESULT_DIR / "figures"
    SUB_DIR = WORK_DIR / "agent_sub_result"
    M4_OUT = WORK_DIR / "m4_output"
    for d in (RESULT_DIR, FIG_DIR, SUB_DIR, M4_OUT):
        d.mkdir(parents=True, exist_ok=True)


# ----------------------- API Keys -----------------------
def set_env(k, v):
    if v:  # 값이 있을 때만 환경변수 설정
        os.environ[k] = v


set_env(
    "OPENAI_API_KEY",
    st.sidebar.text_input(
        "OpenAI API Key",
        type="password",
        value="",  # 항상 공란
        placeholder="sk-****************",
        help="입력 시에만 환경변수를 설정합니다.",
    ),
)
set_env(
    "ENTREZ_EMAIL",
    st.sidebar.text_input(
        "NCBI Entrez Email", value="", placeholder="name@example.com"  # 공란
    ),
)
set_env(
    "CHROMEDRIVER",
    st.sidebar.text_input(
        "ChromeDriver (선택)", value="", placeholder="/path/to/chromedriver"  # 공란
    ),
)

# ----------------------- Paths Summary -----------------------
if WORK_DIR:
    st.sidebar.code(
        f"CODE ROOT: {ROOT}\nWORK DIR:  {WORK_DIR}\n"
        f"- analysis_result/\n- agent_sub_result/\n- m4_output/\n"
    )
else:
    st.sidebar.code(
        f"CODE ROOT: {ROOT}\nWORK DIR:  (미설정)\n"
        f"(좌측에서 작업 폴더를 먼저 적용하세요)\n"
    )

# ----------------------- Module Files (파일 존재만 먼저 표시) -----------------------
FILES = {
    "M1": "A_F_M1_Total_search_1.py",
    "M2": "A_F_M2_Total_deepSearch_4.py",
    "M3": "A_F_M3_TableFigureGen_5.py",
    "M4": "A_F_M4_Gen_article_15_p_to_6-2.py",
    "M5": "A_F_M5_En_3_p.py",
    "M6": "A_F_M6_ReviveScout_1.py",
}

st.sidebar.markdown("**모듈 파일 존재 여부**")
for k, rel in FILES.items():
    ok = (ROOT / rel).exists()
    st.sidebar.markdown(f"- {k} {'✅' if ok else '❌'}")


# ----------------------- (중요) 모듈 로더 — WORK_DIR 설정 후에만 import -----------------------
def load_mod(name, relpath):
    p = ROOT / relpath
    if not p.exists():
        return None
    spec = importlib.util.spec_from_file_location(name, str(p))
    mod = importlib.util.module_from_spec(spec)
    try:
        spec.loader.exec_module(mod)  # type: ignore
    except Exception as e:
        st.sidebar.error(f"{name} 로드 실패: {e}")
        return None
    return mod


mods = {}
if WORK_DIR:
    # WORK_DIR 확정 후에 import (중요!)
    mods = {k: load_mod(k, v) for k, v in FILES.items()}


# ----------------------- Patch module internal paths to WORK_DIR -----------------------
def patch_module_dirs(mods, work_dir: Path):
    # ---- M1: base 디렉토리들을 WORK_DIR로 강제 ----
    if mods.get("M1"):
        try:
            m1 = mods["M1"]
            m1.BASE_DIR = str(work_dir / "agent_sub_result")
            m1.FULLTEXT_DIR = os.path.join(m1.BASE_DIR, "fulltexts")
            m1.INPUT_CSV = os.path.join(m1.BASE_DIR, "pubmed_articles_with_pico.csv")
            m1.EMB_NPY = os.path.join(m1.BASE_DIR, "pico_embeddings.npy")
            m1.OUTPUT_CSV = os.path.join(m1.BASE_DIR, "filtered_articles_semantic.csv")
            os.makedirs(m1.FULLTEXT_DIR, exist_ok=True)
        except Exception as e:
            st.sidebar.warning(f"M1 경로 패치 실패: {e}")

    # ---- M2: WORK_DIR/analysis_result 로 강제 ----
    if mods.get("M2"):
        try:
            mods["M2"].RESULT_DIR = str(work_dir / "analysis_result")
            os.makedirs(mods["M2"].RESULT_DIR, exist_ok=True)
        except Exception as e:
            st.sidebar.warning(f"M2 경로 패치 실패: {e}")

    # ---- M3: base 디렉토리들을 WORK_DIR로 강제 ----
    if mods.get("M3"):
        try:
            m3 = mods["M3"]
            m3.BASE_DIR = str(work_dir / "analysis_result")
            m3.FIG_DIR = os.path.join(m3.BASE_DIR, "figures")
            m3.DISPLAY_CODE_DIR = os.path.join(m3.BASE_DIR, "display_code")
            os.makedirs(m3.FIG_DIR, exist_ok=True)
            os.makedirs(m3.DISPLAY_CODE_DIR, exist_ok=True)
        except Exception as e:
            st.sidebar.warning(f"M3 경로 패치 실패: {e}")


def patch_m6_env(mods, work_dir: Path):
    """M6가 import 직후부터 WORK_DIR를 인지하도록 전역 주입 + 환경변수 설정"""
    m6 = mods.get("M6")
    if not m6:
        return
    try:
        ok = os.getenv("OPENAI_API_KEY", "").strip()
        em = os.getenv("ENTREZ_EMAIL", os.getenv("PUBMED_EMAIL", "")).strip()
        if ok:
            setattr(m6, "OPENAI_API_KEY_HARDCODED", ok)
        if em:
            setattr(m6, "PUBMED_EMAIL_HARDCODED", em)
        # ★ 중요: WORK_DIR 전역 주입 + 환경변수 세팅 (import 시점 로그 방지)
        setattr(m6, "WORK_DIR", str(work_dir))
        os.environ["WORK_DIR"] = str(work_dir)
    except Exception as e:
        st.sidebar.warning(f"M6 env 패치 실패: {e}")


def patch_m6_outdir(mods, work_dir: Path):
    """M6 결과가 항상 WORK_DIR/scout_output/ 로 가도록 ensure_outdirs를 오버라이드"""
    m6 = mods.get("M6")
    if not m6:
        return
    from pathlib import Path as _P

    def _ensure_outdirs_override(_root: _P):
        out = _P(work_dir) / "scout_output"
        out.mkdir(parents=True, exist_ok=True)
        (out / "artifacts").mkdir(exist_ok=True)
        return {"root": out, "arts": out / "artifacts"}

    m6.ensure_outdirs = _ensure_outdirs_override


# WORK_DIR 확정 후에만 패치 호출 (import 직후 바로)
if WORK_DIR and mods:
    patch_module_dirs(mods, WORK_DIR)
    patch_m6_env(mods, WORK_DIR)
    patch_m6_outdir(mods, WORK_DIR)


# ----------------------- Helpers -----------------------
def run(cmd: list[str], input_text: str | None = None, cwd: Path | None = None):
    """서브프로세스 실행 (Windows 콘솔에서도 UTF-8 강제)"""
    try:
        env = os.environ.copy()
        env.setdefault("PYTHONIOENCODING", "utf-8")
        env.setdefault("PYTHONUTF8", "1")
        env.setdefault("LC_ALL", "C.UTF-8")
        env.setdefault("LANG", "C.UTF-8")
        env.setdefault("PYTHONUNBUFFERED", "1")
        r = subprocess.run(
            cmd,
            input=input_text,
            text=True,
            cwd=str(cwd) if cwd else None,
            capture_output=True,
            check=False,
            env=env,
            encoding="utf-8",  # cp949 회피
            errors="replace",  # cp949 회피
        )
        return r.returncode, r.stdout, r.stderr
    except Exception as e:
        return 1, "", f"subprocess error: {e}"


def run_tee(cmd: list[str], cwd: Path | None = None, title: str = "실시간 로그"):
    """자식 프로세스 stdout/stderr를 실시간으로 Streamlit UI와 터미널에 동시에 출력."""
    env = os.environ.copy()
    env.setdefault("PYTHONIOENCODING", "utf-8")
    env.setdefault("PYTHONUTF8", "1")
    env.setdefault("LC_ALL", "C.UTF-8")
    env.setdefault("LANG", "C.UTF-8")
    env.setdefault("PYTHONUNBUFFERED", "1")
    proc = subprocess.Popen(
        cmd,
        cwd=str(cwd) if cwd else None,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
        bufsize=1,  # line-buffered
        env=env,
    )
    q: "queue.Queue[tuple[str,str]]" = queue.Queue()

    def _reader(stream, tag):
        for line in iter(stream.readline, ""):
            q.put((tag, line))
        stream.close()

    t_out = threading.Thread(target=_reader, args=(proc.stdout, "STDOUT"), daemon=True)
    t_err = threading.Thread(target=_reader, args=(proc.stderr, "STDERR"), daemon=True)
    t_out.start()
    t_err.start()
    box = st.empty()
    box.caption(title)
    out_buf: list[str] = []
    err_buf: list[str] = []
    while True:
        try:
            tag, line = q.get(timeout=0.1)
            if tag == "STDOUT":
                out_buf.append(line)
                print(line, end="")
            else:
                err_buf.append(line)
                print(line, end="", file=sys.stderr)
            ui_text = ""
            if out_buf:
                ui_text += "STDOUT\n" + "".join(out_buf[-400:])
            if err_buf:
                ui_text += (
                    ("\n" if ui_text else "") + "STDERR\n" + "".join(err_buf[-200:])
                )
            box.code(ui_text or "(no output yet)")
        except queue.Empty:
            if proc.poll() is not None and q.empty():
                break
    rc = proc.wait()
    return rc, "".join(out_buf), "".join(err_buf)


def show_logs(title, stdout, stderr):
    st.caption(title)
    if stdout:
        st.success("STDOUT")
        st.code(stdout)
    if stderr:
        st.warning("STDERR")
        st.code(stderr)


def validate_docx(path: Path) -> tuple[bool, str]:
    if not path.exists():
        return False, "파일이 존재하지 않습니다."
    if path.suffix.lower() != ".docx":
        return False, "확장자가 .docx 가 아닙니다."
    if path.stat().st_size < 1024:
        return False, f"파일 크기가 너무 작습니다 ({path.stat().st_size} bytes)"
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = set(z.namelist())
            need = {"[Content_Types].xml", "word/document.xml"}
            if not need.issubset(names):
                return False, "DOCX 내부 구조 누락"
    except zipfile.BadZipFile:
        return False, "DOCX(zip) 손상 또는 비정상 파일."
    return True, "정상으로 보입니다."


# ----------------------- Common Guard (도움말은 항상 렌더, 나머지는 안내만) -----------------------
def workdir_ready() -> bool:
    return WORK_DIR is not None


# ----------------------- Tabs -----------------------
t6, t1, t2, t3, t4, t5, t7 = st.tabs(
    [
        "사용법",
        "M1 문헌검색",
        "M2 가설·분석",
        "M3 표/그림 생성",
        "M4 논문(국문)",
        "M5 논문(영문)",
        "M6 아이디어 스카우팅",
    ]
)

# ----------------------- M1 -----------------------
with t1:
    st.header("M1: 문헌 검색 및 분석 보고서 생성")
    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M1"):
            st.error("M1 파일이 없습니다.")
        else:
            m1 = mods["M1"]
            q1 = st.text_input("1) 광범위 키워드")
            q2 = st.text_input("2) 세부 주제(의미검색용)")
            retmax = st.number_input("retmax", 1, 5000, 1000)

            if st.button("실행(M1)"):
                if not (q1 and q2):
                    st.error("모두 입력")
                    st.stop()
                with st.spinner("검색/처리 중..."):
                    ids = m1.search_pubmed(q1, retmax=int(retmax))
                    st.write(f"검색 결과: {len(ids)}편")

                    rows = []
                    for i in range(0, len(ids), 100):
                        for a in m1.fetch_details(ids[i : i + 100]):
                            d = m1.parse_article(a)
                            if d:
                                rows.append(d)
                        time.sleep(0.2)
                    df = pd.DataFrame(rows)
                    (SUB_DIR / "pubmed_articles.csv").write_text(
                        df.to_csv(index=False, encoding="utf-8-sig"), encoding="utf-8"
                    )

                    pico = [m1.extract_pico(x or "") for x in df["Abstract"].fillna("")]
                    df2 = pd.concat([df, pd.DataFrame(pico)], axis=1)
                    (SUB_DIR / "pubmed_articles_with_pico.csv").write_text(
                        df2.to_csv(index=False, encoding="utf-8-sig"), encoding="utf-8"
                    )

                    top = m1.semantic_search(
                        q2,
                        50,
                        20,
                        str(SUB_DIR / "pubmed_articles_with_pico.csv"),
                        str(SUB_DIR / "pico_embeddings.npy"),
                        str(SUB_DIR / "filtered_articles_semantic.csv"),
                    )
                    st.dataframe(top)

                    analyses = {}
                    driver = None
                    try:
                        if os.getenv("CHROMEDRIVER"):
                            driver = m1.setup_driver()
                        for doi in [
                            d for d in top["DOI"].dropna() if str(d).lower() != "na"
                        ]:
                            url = m1.get_paper_url(doi)
                            text = None
                            if url and driver:
                                fp = m1.extract_text_from_url(
                                    driver, url, doi, str(SUB_DIR / "fulltexts")
                                )
                                if fp and Path(fp).exists():
                                    text = Path(fp).read_text(encoding="utf-8")
                            if not text:
                                row = top[top["DOI"] == doi].iloc[0]
                                text = (
                                    row["Abstract"]
                                    if isinstance(row["Abstract"], str)
                                    else ""
                                )
                            if text:
                                analyses[doi] = m1.analyze_paper_with_gpt(text)
                    finally:
                        try:
                            if driver:
                                driver.quit()
                        except Exception:
                            pass

                    rep = m1.consolidate_results(analyses)
                    (SUB_DIR / "consolidated_report.md").write_text(
                        rep, encoding="utf-8"
                    )
                    st.success(f"보고서 저장: {SUB_DIR/'consolidated_report.md'}")
                    st.download_button(
                        "보고서 다운로드", rep.encode("utf-8"), "consolidated_report.md"
                    )

# ----------------------- M2 (Interactive with feedback loops) -----------------------
with t2:
    st.header("M2: 주제 선정 · 데이터 분석")
    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M2"):
            st.error("M2 파일이 없습니다.")
        else:
            m2 = mods["M2"]
            hypothesis = st.text_input("가설")
            method = st.text_input("분석법")
            source = st.radio(
                "데이터 소스", ["파일 업로드", "로컬 경로"], horizontal=True
            )
            uploaded = None
            local_path_str = ""

            if source == "파일 업로드":
                uploaded = st.file_uploader(
                    "CSV/ZIP/GZ/PARQUET (2GB)", type=["csv", "zip", "gz", "parquet"]
                )
            else:
                ph = (
                    str((WORK_DIR / "analysis_result" / "bigdata.csv").resolve())
                    if WORK_DIR
                    else "예: /path/to/bigdata.csv"
                )
                local_path_str = st.text_input("서버/로컬 경로", placeholder=ph)

            def _resolve_src():
                csv_path = RESULT_DIR / "uploaded.csv"
                if source == "파일 업로드":
                    if not uploaded:
                        st.error("파일 선택")
                        st.stop()
                    suffix = uploaded.name.split(".")[-1].lower()
                    tmp_path = RESULT_DIR / ("_incoming_" + uploaded.name)
                    with open(tmp_path, "wb") as f:
                        shutil.copyfileobj(uploaded, f, length=1024 * 1024)
                    if suffix == "csv":
                        return tmp_path
                    if suffix == "gz":
                        with gzip.open(tmp_path, "rb") as gz, open(
                            csv_path, "wb"
                        ) as out:
                            shutil.copyfileobj(gz, out)
                        return csv_path
                    if suffix == "zip":
                        with zipfile.ZipFile(tmp_path) as z:
                            cvs = [
                                m for m in z.namelist() if m.lower().endswith(".csv")
                            ]
                            if not cvs:
                                st.error("ZIP 내 CSV 없음")
                                st.stop()
                            z.extract(cvs[0], RESULT_DIR)
                            return RESULT_DIR / cvs[0]
                    if suffix == "parquet":
                        pd.read_parquet(tmp_path).to_csv(csv_path, index=False)
                        return csv_path
                    st.error("확장자 미지원")
                    st.stop()
                else:
                    p = Path(local_path_str)
                    if not p.exists():
                        st.error("경로 없음")
                        st.stop()
                    if p.suffix.lower() == ".parquet":
                        pd.read_parquet(p).to_csv(csv_path, index=False)
                        return csv_path
                    if p.suffix.lower() == ".gz":
                        with gzip.open(p, "rb") as gz, open(csv_path, "wb") as out:
                            shutil.copyfileobj(gz, out)
                            return csv_path
                    if p.suffix.lower() == ".zip":
                        with zipfile.ZipFile(p) as z:
                            cvs = [
                                m for m in z.namelist() if m.lower().endswith(".csv")
                            ]
                            if not cvs:
                                st.error("ZIP 내 CSV 없음")
                                st.stop()
                            z.extract(cvs[0], RESULT_DIR)
                            return RESULT_DIR / cvs[0]
                    return p

            # ---------------- M2 단계 시작 ----------------
            if st.button("M2 시작"):
                if not (hypothesis and method):
                    st.error("가설/분석법 입력")
                    st.stop()
                st.session_state.m2_src = str(_resolve_src())
                st.session_state.m2_step = 1
                st.session_state.m2_summary = None
                st.session_state.m2_hv = ""
                st.session_state.m2_hv_fb = ""
                st.session_state.m2_code = ""
                st.session_state.m2_code_fb = ""
                st.session_state.m2_run_out = ""
                st.rerun()

            if st.session_state.get("m2_step"):
                step = st.session_state.m2_step
                src = st.session_state.m2_src

                # ---------------- step 1 ----------------
                if step == 1:
                    st.subheader("1) CSV 로드/요약")
                    df, summary = m2.analyze_csv(src)
                    if df is None:
                        st.error(summary)
                        st.stop()
                    st.session_state.m2_summary = summary
                    st.json(summary)
                    if st.button("확인(다음) → 가설 검증"):
                        st.session_state.m2_step = 2
                        st.rerun()

                # ---------------- step 2 ----------------
                elif step == 2:
                    st.subheader("2) 가설 검증 → 피드백 루프")

                    # 처음 한 번만 검증 실행
                    if not st.session_state.get("m2_hv"):
                        st.session_state.m2_hv = m2.validate_hypothesis(
                            hypothesis, method, st.session_state.m2_summary, src
                        )

                    # 결과 표시
                    st.text_area(
                        "가설 검증 결과",
                        st.session_state.m2_hv,
                        height=200,
                        key="m2_hv_view",
                    )
                    st.session_state.m2_hv_fb = st.text_area(
                        "피드백 (수정 요청/보완점)",
                        st.session_state.get("m2_hv_fb", ""),
                        height=120,
                    )

                    c1, c2, c3 = st.columns(3)
                    if c1.button("피드백 반영하여 재검증", key="btn_hv_retest"):
                        fb_text = (st.session_state.get("m2_hv_fb") or "").strip()
                        st.session_state.m2_hv = m2.validate_hypothesis(
                            hypothesis
                            + ("\n[Feedback]: " + fb_text if fb_text else ""),
                            method,
                            st.session_state.m2_summary,
                            src,
                        )
                        st.toast("재검증 완료")
                        st.rerun()

                    if c2.button("확정(다음) → 코드 생성"):
                        st.session_state.m2_step = 3
                        st.rerun()

                    if c3.button("중단"):
                        del st.session_state["m2_step"]
                        st.rerun()

                # ---------------- step 3 ----------------
                elif step == 3:
                    st.subheader("3) 코드 생성 → 피드백 루프")
                    prompt, code = m2.generate_analysis_code(
                        hypothesis,
                        method,
                        st.session_state.m2_summary,
                        src,
                        st.session_state.m2_hv,
                    )
                    st.session_state.m2_code = st.text_area(
                        "생성 코드(수정 가능)", code, height=320
                    )
                    st.session_state.m2_code_fb = st.text_area(
                        "피드백 (변수명/시각화/모형 등)",
                        st.session_state.get("m2_code_fb", ""),
                        height=120,
                    )

                    c1, c2, c3 = st.columns(3)
                    if c1.button("피드백 반영 재생성", key="btn_code_regen"):
                        hypo_with_fb = hypothesis + (
                            "\n[Developer feedback]:\n" + st.session_state.m2_code_fb
                            if st.session_state.get("m2_code_fb")
                            else ""
                        )
                        _, code2 = m2.generate_analysis_code(
                            hypo_with_fb,
                            method,
                            st.session_state.m2_summary,
                            src,
                            st.session_state.m2_hv,
                        )
                        st.session_state.m2_code = code2
                        st.toast("재생성 완료")
                        st.rerun()

                    if c2.button("확정(다음) → 실행"):
                        st.session_state.m2_step = 4
                        st.rerun()

                    if c3.button("중단"):
                        del st.session_state["m2_step"]
                        st.rerun()

                # ---------------- step 4 ----------------
                elif step == 4:
                    st.subheader("4) 실행 → Results/Discussion 편집")
                    out = m2.save_and_run_code(
                        st.session_state.m2_code,
                        "",
                        str(RESULT_DIR / "final_analysis.py"),
                        df=None,
                        summary=st.session_state.m2_summary,
                        input_type="CSV",
                    )
                    st.session_state.m2_run_out = out or ""
                    st.code(st.session_state.m2_run_out or "(no output)")

                    res = m2.gpt_analysis_results(st.session_state.m2_run_out or "")
                    disc = m2.gpt_discussion(hypothesis, res)
                    res_e = st.text_area("Results(수정)", res, height=200)
                    disc_e = st.text_area("Discussion(수정)", disc, height=200)
                    if st.button("최종 저장"):
                        final_md = (
                            "## 가설 검증(최종)\n"
                            + st.session_state.m2_hv
                            + "\n\n## Results\n"
                            + res_e
                            + "\n\n## Discussion\n"
                            + disc_e
                        )
                        (RESULT_DIR / "final_analysis.md").write_text(
                            final_md, encoding="utf-8"
                        )
                        (RESULT_DIR / "final_analysis.json").write_text(
                            json.dumps(
                                {
                                    "HypothesisValidationFinal": st.session_state.m2_hv,
                                    "Results": res_e,
                                    "Discussion": disc_e,
                                },
                                ensure_ascii=False,
                                indent=2,
                            ),
                            encoding="utf-8",
                        )
                        st.success(f"저장 완료: {RESULT_DIR/'final_analysis.md'}")
                        del st.session_state["m2_step"]


# ----------------------- M3 (Bulk list edit → refine → run all or single) -----------------------
with t3:
    st.header("M3: 분석 자료 시각화")
    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M3"):
            st.error("M3 파일이 없습니다.")
        else:
            m3 = mods["M3"]

            # 0) 초기 추천 불러오기
            if st.button("추천 불러오기"):
                final_md = (
                    (RESULT_DIR / "final_analysis.md").read_text(encoding="utf-8")
                    if (RESULT_DIR / "final_analysis.md").exists()
                    else ""
                )
                code_py = (
                    (RESULT_DIR / "final_analysis.py").read_text(encoding="utf-8")
                    if (RESULT_DIR / "final_analysis.py").exists()
                    else ""
                )
                if not (final_md and code_py):
                    st.error("M2 산출물(final_analysis.md/py)이 필요합니다.")
                else:
                    try:
                        raw = m3.ask_items(final_md, code_py, feedback="")
                    except Exception as e:
                        st.error(f"추천 생성 실패: {e}")
                        st.stop()
                    st.session_state.m3_items = m3.parse_items(raw) or []
                    st.session_state.m3_base = code_py
                    st.toast("추천 항목을 불러왔습니다.")

            # Helpers
            def _items_to_df(items: list[dict]) -> pd.DataFrame:
                if not items:
                    return pd.DataFrame(
                        columns=["display_type", "desc", "purpose", "feedback"]
                    )
                return pd.DataFrame(
                    [
                        {
                            "display_type": it.get("display_type", ""),
                            "desc": it.get("desc", ""),
                            "purpose": it.get("purpose", ""),  # <-- 추가
                            "feedback": it.get("feedback", ""),
                        }
                        for it in items
                    ]
                )

            def _df_to_items(df: pd.DataFrame) -> list[dict]:
                items = []
                for _, row in df.iterrows():
                    dt = str(row.get("display_type", "")).strip()
                    dc = str(row.get("desc", "")).strip()
                    pu = str(row.get("purpose", "")).strip()  # <-- 추가
                    fb = str(row.get("feedback", "")).strip()
                    if dt or dc:
                        items.append(
                            {
                                "display_type": dt,
                                "desc": dc,
                                "purpose": pu,
                                "feedback": fb,
                            }
                        )
                return items

            if st.session_state.get("m3_items") is not None:
                items: list[dict] = st.session_state.m3_items
                st.subheader(
                    f"추천 목록 ({len(items)}개) — 아래에서 **추가/수정/삭제** 후 적용"
                )

                df_src = _items_to_df(items)
                df_src["삭제"] = False
                edited = st.data_editor(
                    df_src,
                    num_rows="dynamic",
                    use_container_width=True,
                    column_config={
                        "display_type": st.column_config.TextColumn("display_type"),
                        "desc": st.column_config.TextColumn("desc", width="large"),
                        "purpose": st.column_config.TextColumn(
                            "purpose", width="xlarge"
                        ),  # <-- 추가
                        "feedback": st.column_config.TextColumn(
                            "feedback", width="medium"
                        ),
                        "삭제": st.column_config.CheckboxColumn("삭제"),
                    },
                    key="m3_editor",
                )
                c_apply, c_reset = st.columns(2)
                if c_apply.button("변경사항 적용(삭제 포함)"):
                    keep = edited[~edited["삭제"]].drop(
                        columns=["삭제"], errors="ignore"
                    )
                    st.session_state.m3_items = _df_to_items(keep)
                    st.success(f"목록 적용 완료: {len(st.session_state.m3_items)}개")
                    st.rerun()
                if c_reset.button("목록 되돌리기"):
                    st.info("편집 전 상태로 유지했습니다.")

                st.divider()

                st.subheader("전역 피드백으로 재추천 (선택)")
                global_fb = st.text_area(
                    "이 피드백을 반영해 '전체 목록'을 재작성합니다. (비워두면 현재 목록 유지)",
                    value="",
                    height=120,
                    placeholder="예) 표1: 기초통계, 그림1: ROC, 그림2: SHAP... 색상 팔레트 통일",
                )
                if st.button("전역 피드백 반영 → 새 목록 받기"):
                    final_md = (
                        (RESULT_DIR / "final_analysis.md").read_text(encoding="utf-8")
                        if (RESULT_DIR / "final_analysis.md").exists()
                        else ""
                    )
                    code_py = (
                        (RESULT_DIR / "final_analysis.py").read_text(encoding="utf-8")
                        if (RESULT_DIR / "final_analysis.py").exists()
                        else ""
                    )
                    if not (final_md and code_py):
                        st.error("M2 산출물(final_analysis.md/py)이 필요합니다.")
                    else:
                        prompt = (
                            "Rewrite the WHOLE list incorporating ONLY the feedback below.\n"
                            "OUTPUT STRICTLY one item per line, numbered, in this exact format (no bullets/JSON/code):\n"
                            "1. <display_type> - <desc>\n"
                            "2. <display_type> - <desc>\n"
                            "Use only variables/measures that appear in final_analysis.md.\n"
                            "No extra text, no headings, no explanations.\n"
                            f"[User feedback]\n{global_fb}\n"
                            f"[Current items]\n{st.session_state.m3_items}"
                        )
                        raw2 = m3.ask_items(final_md, code_py, prompt)
                        new_items = m3.parse_items(raw2) or []
                        if new_items:
                            st.session_state.m3_items = new_items
                            st.success(f"재추천 성공: {len(new_items)}개로 업데이트")
                        else:
                            st.warning("재추천 결과가 비어있습니다. 기존 목록 유지")

                st.divider()

                # 실행
                st.subheader("실행")
                base_code = st.session_state.get("m3_base", "")

                def _run_item(it: dict) -> bool:
                    try:
                        # 🔒 앱 레벨에서도 저장 강제 규칙을 피드백에 주입
                        forced_rule = (
                            '플롯 마지막에는 반드시 m_tidy_finalize('
                            'name="Figure_1", fig=fig, ax=ax, legend=leg, width="double")를 '
                            '호출해 저장까지 완료하세요.'
                        )
                        merged_feedback = ((it.get("feedback", "") or "").strip() + "\n" + forced_rule).strip()

                        # generate_code(item, code_text, feedback)
                        code = m3.generate_code(it, base_code, merged_feedback)
                    except Exception as e:
                        st.error(f"코드 생성 실패: {e}")
                        return False
                    try:
                        return m3.run_with_retry(
                            code,
                            f"final_with_{it.get('display_type','item')}.py",
                            base_code,
                            max_retries=3,
                        )
                    except Exception as e:
                        st.error(f"실행 실패: {e}")
                        return False

                if st.session_state.m3_items:
                    idx_options = list(range(len(st.session_state.m3_items)))
                    sel_idx = st.selectbox(
                        "단건 실행할 항목 선택",
                        idx_options,
                        # ✅ 괄호 누락 수정 + purpose 표시 유지
                        format_func=lambda i: (
                            f"{i+1}. "
                            f"{st.session_state.m3_items[i].get('display_type','?')} – "
                            f"{st.session_state.m3_items[i].get('desc','')[:40]}  "
                            f"({st.session_state.m3_items[i].get('purpose','')[:22]})"
                        ),
                    )
                    c_one, c_all = st.columns(2)

                    if c_one.button("이 항목만 생성"):
                        it = st.session_state.m3_items[sel_idx]
                        ok = _run_item(it)
                        st.toast("성공" if ok else "실패")

                    # ✅ disabled 인자 위치/괄호 오류 수정
                    busy = st.session_state.get("m3_busy", False)    
                    if c_all.button("연속 생성(전체)", disabled=busy):
                        st.session_state.m3_busy = True
                        try:
                            items_to_run = st.session_state.m3_items
                            st.info(f"총 {len(items_to_run)}개 항목을 한 번에 실행합니다…")
                            ok = m3.run_batch(
                                items_to_run,
                                base_code,
                                filename="final_batch_run.py",
                                max_retries=5,
                            )
                            if ok:
                                st.success(f"배치 실행 완료 ✅ 결과는 {(RESULT_DIR / 'figures')} 에 저장되었습니다.")
                            else:
                                st.error("배치 실행 중 일부 실패가 발생했습니다.")
                        finally:
                            st.session_state.m3_busy = False
                else:
                    st.info("실행할 항목이 없습니다. 먼저 목록을 구성하세요.")

# ----------------------- M4 (section-by-section loop) -----------------------
with t4:
    st.header("M4: 초안 논문 생성기")
    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M4"):
            st.error("M4 파일이 없습니다.")
        else:
            # ✅ 단어 수 입력창 제거 (SECTION_MIN은 M4 모듈 내부에서 자동 적용)
            max_tokens = st.number_input("Max tokens", 1000, 12000, 7000, step=500)
            temperature = st.slider("Temperature", 0.0, 1.0, 0.3, 0.05)

            # ─────────────────────────────────────────────────────────────
            # Helper: DOCX → HTML 미리보기 (mammoth 사용, 없으면 안내)
            # ─────────────────────────────────────────────────────────────
            def _docx_to_html(docx_path: Path) -> str:
                try:
                    import mammoth, base64
                except Exception:
                    return "__MAMMOTH_MISSING__"

                def _image_handler(image):
                    with image.open() as f:
                        data = f.read()
                    b64 = base64.b64encode(data).decode("utf-8")
                    ext = image.content_type or "image/png"
                    return {"src": f"data:{ext};base64,{b64}"}

                with open(docx_path, "rb") as f:
                    res = mammoth.convert_to_html(
                        f, convert_image=mammoth.images.inline(_image_handler)
                    )
                css = """
                <style>
                body{font-family:-apple-system, Segoe UI, Roboto, Noto Sans KR, sans-serif; line-height:1.6;}
                h1,h2,h3{margin:1.2em 0 .4em;}
                p{margin:.5em 0;}
                figure{margin:1em 0; text-align:center;}
                figcaption{font-size:.9em; color:#666;}
                table{border-collapse:collapse; width:100%; margin:.6em 0;}
                th,td{border:1px solid #ddd; padding:6px 8px;}
                </style>
                """
                return css + res.value

            # ─────────────────────────────────────────────────────────────
            # 실행 버튼 (원클릭 생성)
            # ─────────────────────────────────────────────────────────────
            if st.button("M4 실행(원클릭)"):
                with st.spinner("생성 중..."):
                    try:
                        out_path = mods["M4"].generate_manuscript(
                            project=str(WORK_DIR),
                            # min_words 없음 (내부 SECTION_MIN 사용)
                            max_tokens=int(max_tokens),
                            temperature=float(temperature),
                            output_path=None,  # 기본: WORK_DIR/m4_output/manuscript.docx
                        )
                        st.success(f"완료: {out_path}")
                        st.download_button(
                            "manuscript_ko.docx",
                            Path(out_path).read_bytes(),
                            "manuscript_ko.docx",
                        )
                    except Exception as e:
                        st.error(f"오류: {e}")

            # 상태파일이 있으면 섹션 편집 + 자동개정 + 미리보기 제공
            state_p = M4_OUT / "m4_state.json"
            if state_p.exists():
                st.info(
                    "섹션별 텍스트가 생성되었습니다. 아래에서 직접 수정하거나, 지시문으로 자동 개정한 뒤 재빌드하세요."
                )
                state = json.loads(state_p.read_text(encoding="utf-8"))

                order = [
                    "abstract",
                    "introduction",
                    "methods",
                    "results",
                    "discussion",
                    "conclusion",
                    "references",
                ]

                # 섹션 편집 루프
                for sec in order:
                    st.markdown(f"### {sec}")
                    txt = state.get("sections", {}).get(sec, {}).get("text", "") or ""
                    edited = st.text_area(
                        f"{sec} (editable)", txt, height=220, key=f"m4_txt_{sec}"
                    )
                    # 즉시 저장(수동)
                    if st.button(f"{sec} 저장", key=f"m4_save_{sec}"):
                        state.setdefault("sections", {}).setdefault(sec, {})[
                            "text"
                        ] = edited
                        state_p.write_text(
                            json.dumps(state, ensure_ascii=False, indent=2),
                            encoding="utf-8",
                        )
                        st.toast("저장 완료")

                    # ✅ 자동 개정: 지시문 입력 → LLM으로 해당 섹션만 고쳐쓰기
                    with st.expander(
                        f"{sec} · 자동 개정 (지시문 입력)", expanded=False
                    ):
                        fb_auto = st.text_area(
                            "지시문 (예: '두괄식 정리하고 수치·근거를 더 구체화')",
                            value="",
                            height=110,
                            key=f"m4_auto_fb_{sec}",
                        )
                        c1, c2 = st.columns(2)
                        if c1.button(
                            f"지시문 반영하여 자동 수정 · {sec}",
                            key=f"m4_btn_autorevise_{sec}",
                        ):
                            try:
                                # M4 모듈의 client/model/SECTION_MIN 사용
                                client = mods["M4"].make_client()
                                model = (state.get("options", {}) or {}).get(
                                    "model"
                                ) or mods["M4"].CFG["model"]
                                target_min = max(
                                    mods["M4"].SECTION_MIN.get(sec, 500), 500
                                )

                                revised = (
                                    mods["M4"]
                                    ._ask(
                                        client,
                                        model,
                                        "섹션 개정: 피드백을 반영해 자연스럽게 고쳐쓰기. 중복 최소화, 사실 기반.",
                                        [
                                            {
                                                "type": "input_text",
                                                "text": f"[현재 초안]\n{edited[-6000:]}",
                                            },
                                            {
                                                "type": "input_text",
                                                "text": f"[피드백]\n{fb_auto}",
                                            },
                                            {
                                                "type": "input_text",
                                                "text": f"[제약] 최소 {target_min} 단어, 목록/표/그림 금지(결과 제외), 한국어 학술문체",
                                            },
                                        ],
                                        temp=float(temperature),
                                        max_tokens=min(3000, int(max_tokens)),
                                    )
                                    .strip()
                                )

                                if revised:
                                    state.setdefault("sections", {}).setdefault(
                                        sec, {}
                                    )["text"] = revised
                                    state_p.write_text(
                                        json.dumps(state, ensure_ascii=False, indent=2),
                                        encoding="utf-8",
                                    )
                                    st.success("자동 수정 완료! 본문에 반영되었습니다.")
                                    st.experimental_rerun()
                                else:
                                    st.warning(
                                        "개정 결과가 비었습니다. 지시문을 더 구체화해보세요."
                                    )
                            except Exception as e:
                                st.error(f"자동 수정 실패: {e}")

                        if c2.button(
                            f"{sec} 취소 (원문으로 되돌리기)",
                            key=f"m4_btn_cancel_{sec}",
                        ):
                            # 상태파일에서 다시 읽어 화면만 갱신
                            state = json.loads(state_p.read_text(encoding="utf-8"))
                            st.experimental_rerun()

                st.divider()

                # ✅ 라이브 미리보기 (DOCX 없이 텍스트 바로 렌더)
                st.subheader("라이브 미리보기 (DOCX 없이)")
                ko = {
                    "abstract": "초록",
                    "introduction": "서론",
                    "methods": "방법",
                    "results": "결과",
                    "discussion": "고찰",
                    "conclusion": "결론",
                    "references": "참고문헌",
                }
                title_txt = (
                    state.get("title", "")
                    or state.get("title_en", "")
                    or "미리보기 제목"
                )
                buf = [f"# {title_txt}\n"]
                import re as _re

                for k in order:
                    t = (
                        state.get("sections", {}).get(k, {}).get("text", "") or ""
                    ).strip()
                    if not t:
                        continue
                    buf.append(f"## {ko[k]}")
                    for p in _re.split(r"\n\s*\n", t):
                        if p.strip():
                            buf.append(p.strip())
                st.markdown("\n\n".join(buf))

                st.divider()

                # ✅ DOCX 재빌드 (LLM 호출 없음) + DOCX→HTML 미리보기
                if st.button("수정 반영하여 DOCX 재빌드"):
                    try:
                        new_out = mods["M4"].build_docx(str(WORK_DIR), state, None)
                        st.success(f"재빌드 완료: {new_out}")
                        st.download_button(
                            "manuscript_ko.docx (rebuild)",
                            Path(new_out).read_bytes(),
                            "manuscript_ko.docx",
                        )
                    except Exception as e:
                        st.error(f"재빌드 오류: {e}")

                docx_path = M4_OUT / "manuscript.docx"
                if docx_path.exists():
                    if st.button("DOCX 미리보기 열기"):
                        html_str = _docx_to_html(docx_path)
                        if html_str == "__MAMMOTH_MISSING__":
                            st.warning(
                                "mammoth가 설치되어 있지 않습니다. `pip install mammoth` 후 다시 시도하세요."
                            )
                        else:
                            from streamlit.components.v1 import html as st_html

                            st_html(html_str, height=800, scrolling=True)
                else:
                    st.info(
                        "아직 DOCX가 없습니다. 먼저 ‘M4 실행(원클릭)’ 또는 ‘재빌드’를 해주세요."
                    )
            else:
                st.info(
                    "아직 M4 상태파일(m4_state.json)이 없습니다. 먼저 ‘M4 실행(원클릭)’을 실행하세요."
                )

# ----------------------- M5 (영문 논문 섹션별 피드백 루프) -----------------------
with t5:
    st.header("M5: 영문 논문 (섹션별 피드백 루프)")
    st.info("💡 M3를 먼저 실행하면 생성된 `R3_figures`가 M5 문서에 자동 삽입됩니다.")

    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M5"):
            st.error("M5 파일이 없습니다.")
        else:
            # 안전하게 m4_output 경로 계산 (M4_OUT이 None일 수 있으므로 직접 계산)
            m4_out_dir = Path(WORK_DIR) / "m4_output"
            state_default = m4_out_dir / "m4_state_en.json"

            locale = st.selectbox("APA Locale", ["en-US", "en-GB"], 0)
            non = st.checkbox(
                "초기 초안은 --non-interactive로 생성", True, key="m5_non_init"
            )

            if st.button("M5 초기 초안 생성"):
                cmd = [
                    sys.executable,
                    FILES["M5"],
                    "--project",
                    str(WORK_DIR),
                    "--locale",
                    locale,
                ]
                if non:
                    cmd.append("--non-interactive")
                rc, out, err = run_tee(
                    cmd, title="M5 실시간 로그"
                )  # 실시간 + 터미널 동시 출력
                show_logs("M5 로그 (요약)", out, err)

                # 상태파일이 생겼으면 초기 인덱스와 경로 저장
                if state_default.exists():
                    st.session_state.m5_idx = 0
                    st.session_state.m5_state_path = str(state_default)

            # 상태파일 존재 시 섹션별 편집 UI 노출
            state_path = Path(st.session_state.get("m5_state_path", state_default))
            if state_path.exists():
                state = json.loads(state_path.read_text(encoding="utf-8"))
                order = [
                    "abstract",
                    "introduction",
                    "methods",
                    "results",
                    "discussion",
                    "conclusion",
                    "references",
                ]

                idx = st.session_state.get("m5_idx", 0)
                idx = max(0, min(idx, len(order) - 1))
                sec = order[idx]

                st.subheader(f"{idx+1}/{len(order)} · {sec}")
                text = state.get("sections_en", {}).get(sec, {}).get("text", "") or ""
                fb = st.text_area(
                    "Feedback",
                    state.get("sections_en", {}).get(sec, {}).get("feedback", "") or "",
                    height=120,
                    key=f"m5_fb_{sec}",
                )
                edited = st.text_area(
                    "Section text (editable)", text, height=300, key=f"m5_txt_{sec}"
                )

                c1, c2, c3, c4 = st.columns(4)
                if c1.button("Save feedback"):
                    state.setdefault("sections_en", {}).setdefault(sec, {})[
                        "feedback"
                    ] = fb
                    state_path.write_text(
                        json.dumps(state, ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
                    st.toast("Saved")
                if c2.button("Save text (confirm)"):
                    state.setdefault("sections_en", {}).setdefault(sec, {})[
                        "text"
                    ] = edited
                    state_path.write_text(
                        json.dumps(state, ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
                    st.toast("Saved")
                if c3.button("Prev"):
                    st.session_state.m5_idx = max(0, idx - 1)
                    st.rerun()
                if c4.button("Next"):
                    st.session_state.m5_idx = min(len(order) - 1, idx + 1)
                    st.rerun()

                if st.button("Rebuild DOCX (EN)"):
                    doc = Document()
                    style = doc.styles["Normal"]
                    style.font.name = "Times New Roman"
                    style.font.size = Pt(11)

                    def H(t):
                        doc.add_heading(t, 1).alignment = WD_ALIGN_PARAGRAPH.LEFT

                    H("Title")
                    doc.add_paragraph(state.get("title_en", "") or "")
                    for k, h in [
                        ("abstract", "Abstract"),
                        ("introduction", "Introduction"),
                        ("methods", "Methods"),
                        ("results", "Results"),
                        ("discussion", "Discussion"),
                        ("conclusion", "Conclusion"),
                        ("references", "References"),
                    ]:
                        H(h)
                        t = (
                            state.get("sections_en", {}).get(k, {}).get("text", "")
                            or ""
                        ).strip()
                        for p in re.split(r"\n\s*\n", t):
                            if p.strip():
                                doc.add_paragraph(p.strip())

                    out = m4_out_dir / "manuscript_en.docx"
                    m4_out_dir.mkdir(parents=True, exist_ok=True)
                    doc.save(out)
                    ok, msg = validate_docx(out)
                    st.write(f"DOCX 검사: {msg}")
                    st.download_button(
                        "manuscript_en.docx", out.read_bytes(), "manuscript_en.docx"
                    )
            else:
                st.info(
                    "아직 영어 상태파일(m4_state_en.json)이 없습니다. 먼저 ‘M5 초기 초안 생성’을 실행하세요."
                )

# ----------------------- Help -----------------------
with t6:
    # ---- Fonts & CSS ----
    st.markdown(
        """
    <link href="https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;500;700;900&family=JetBrains+Mono:wght@400;600&display=swap" rel="stylesheet">
    <style>
    :root{
        --brand:#4F46E5; /* 인디고 */
        --brand-2:#06B6D4; /* 청록 */
        --ink:#0F172A;     /* 거의 검정 */
        --muted:#475569;   /* 잿빛 */
        --bg:#0B1020;      /* 히어로 배경 */
    }
    html, body, [class*="css"]  {
        font-family: 'Noto Sans KR', system-ui, -apple-system, Segoe UI, Roboto, 'Apple SD Gothic Neo', 'Malgun Gothic', sans-serif;
        color: var(--ink);
    }
    .block-container {max-width: 1100px;}
    .hero{
        background: linear-gradient(135deg, rgba(79,70,229,.12), rgba(6,182,212,.12));
        border: 1px solid rgba(79,70,229,.25);
        padding: 28px 28px;
        border-radius: 18px;
        margin-bottom: 18px;
    }
    .hero-title{
        font-weight: 900; font-size: 28px; letter-spacing:-.2px;
        margin:0 0 10px 0;
    }
    .hero-sub{color: var(--muted); margin:0;}
    .section-title{
        font-weight: 800; font-size: 20px; margin: 24px 0 10px 0;
    }
    .ol{
        counter-reset: num;
        margin: 0; padding-left: 0;
    }
    .ol li{
        list-style: none;
        margin: 8px 0; padding-left: 44px; position: relative;
    }
    .ol li::before{
        counter-increment: num;
        content: counter(num);
        position: absolute; left: 0; top: 0;
        width: 32px; height: 32px; line-height: 32px; text-align:center;
        border-radius: 10px;
        background: linear-gradient(135deg, var(--brand), var(--brand-2));
        color: white; font-weight: 700;
    }
    .cards{display: grid; grid-template-columns: repeat(2, minmax(0,1fr)); gap: 14px;}
    @media (max-width: 900px){ .cards{ grid-template-columns: 1fr; } }
    .card{
        border: 1px solid rgba(15,23,42,.12);
        border-radius: 16px; padding: 16px;
        background: white; transition: transform .12s ease, box-shadow .12s ease;
    }
    .card:hover{ transform: translateY(-2px); box-shadow: 0 6px 18px rgba(15,23,42,.08); }
    .card h4{ margin:0 0 6px 0; font-weight: 800; }
    .hint{ color: var(--muted); font-size: 14px; margin-top:6px;}
    .mono{ font-family: 'JetBrains Mono', ui-monospace, SFMono-Regular, Menlo, Consolas, monospace; font-size: 13px;}
    .pill{
        display:inline-block; font-size:12px; font-weight:700; padding:4px 10px; border-radius:999px;
        background: rgba(79,70,229,.12); color: #3730A3; margin-left: 6px; vertical-align: middle;
    }
    </style>
    """,
        unsafe_allow_html=True,
    )

    # ---- Hero ----
    st.markdown(
        f"""
    <div class="hero">
    <div class="hero-title">안녕하세요, 인하대병원 의생명연구원님 <span class="pill">AGENT&nbsp;AI</span></div>
    <p class="hero-sub">연구 지원을 위한 GPT 기반 자동화 시스템입니다.</p>
    </div>
    """,
        unsafe_allow_html=True,
    )

    # ---- 기초 설정 ----
    st.markdown('<div class="section-title">기초 설정</div>', unsafe_allow_html=True)
    st.markdown(
        """
    <ol class="ol">
    <li>좌측 <b>환경 설정</b>에서 출력물을 저장할 <b>폴더 경로</b>를 입력하세요.</li>
    <li>본 시스템은 <b>OpenAI GPT</b> 기반으로 동작합니다. <br>
        보유하신 <b>OpenAI API Key</b>를 입력하세요. <span class="mono">sk-...</span></li>
    <li>문헌 검색을 위해 <b>이메일 주소</b>를 입력하세요.</li>
    <li>정밀한 문헌 크롤링을 위해 <b>ChromeDriver 경로</b>를 입력하세요.</li>
    </ol>
    """,
        unsafe_allow_html=True,
    )

    # ---- LLM AGENT 모듈 ----
    st.markdown(
        '<div class="section-title">LLM AGENT 모듈</div>', unsafe_allow_html=True
    )
    st.markdown(
        """
    <div class="cards">
    <div class="card">
        <h4>🔎 M1 · 문헌 조사</h4>
        원하는 키워드에 대한 최신/핵심 문헌을 자동 수집·정리합니다.
        <div class="hint">검색 전략, 필터, 요약 PDF까지 생성</div>
    </div>
    <div class="card">
        <h4>🧪 M2 · 데이터 분석</h4>
        수집된 결과를 바탕으로 <b>보유 데이터</b>를 자동 전처리/분석합니다.
        <div class="hint">통계 검정 · 피처 엔지니어링 · 재현성 있는 파이프라인</div>
    </div>
    <div class="card">
        <h4>📊 M3 · 시각화</h4>
        분석 결과를 논문급 도해로 변환합니다.
        <div class="hint">Figure 레이아웃 · 캡션 · 해상도 자동화</div>
    </div>
    <div class="card">
        <h4>📝 M4 · 1차 문헌 생성</h4>
        선행 조사·분석·시각화 기반으로 초고를 생성합니다.
        <div class="hint">IMRaD 구조 · 레퍼런스 자리표시자 포함</div>
    </div>
    <div class="card">
        <h4>🧷 M5 · 논문 형식 교정</h4>
        저널 가이드라인에 맞춰 포맷/문체/참고문헌을 교정합니다.
        <div class="hint">저널별 템플릿 · 표/그림 번호 매기기 · 인용 일관성</div>
    </div>
    </div>
    """,
        unsafe_allow_html=True,
    )

# ----------------------- M6 (ReviveScout: 초안폴더 → 신규 주제 스카우팅) -----------------------
with t7:
    st.header("M6: 기존 작성 중단논문 기반 신규 주제 스카우팅")
    if not workdir_ready():
        st.info("⬅️ 좌측에서 ‘작업 폴더 적용’을 먼저 진행해주세요.")
    else:
        if not mods.get("M6"):
            st.error(
                "M6 파일이 없습니다. A_F_M6_ReviveScout_1.py 를 코드 폴더에 두세요."
            )
        else:
            draft_root = st.text_input(
                "초안 폴더(.docx/.hwp/.hwpx가 들어있는 폴더)",
                value=(str(WORK_DIR) if WORK_DIR else ""),
                placeholder="예: /path/to/drafts",
            )
            years = st.number_input(
                "최근 N년 스카우팅", min_value=2, max_value=15, value=5, step=1
            )
            max_cands = st.number_input(
                "최대 후보 토픽 수", min_value=6, max_value=50, value=12, step=1
            )
            m6_model = st.selectbox(
                "LLM 모델(권장: gpt-4o-mini)",
                ["gpt-4o-mini", "gpt-4o", "o3-mini"],
                index=0,
            )

            colr1, colr2 = st.columns(2)
            if colr1.button("실행(M6)", use_container_width=True):
                p = Path(draft_root)
                if not p.exists():
                    st.error("초안 폴더 경로가 존재하지 않습니다.")
                    st.stop()
                cmd = [
                    sys.executable,
                    "-u",
                    FILES["M6"],
                    "--folder",
                    str(p.resolve()),
                    "--years",
                    str(int(years)),
                    "--max-candidates",
                    str(int(max_cands)),
                    "--model",
                    m6_model,
                ]
                with st.spinner("ReviveScout 실행 중…"):
                    # 완전 무출력 실행 (stdout/stderr 모두 무시)
                    rc, _, _ = run(cmd, cwd=ROOT)
                if rc == 0:
                    st.success("M6 완료 ✅")
                else:
                    st.warning(
                        "M6가 정상 종료되지 않았습니다. (자세한 로그는 표시하지 않습니다)"
                    )

            # 산출물 미리보기 / 다운로드
            st.divider()
            st.subheader("산출물 (scout_output/)")
            scout_dir = Path(draft_root).resolve() / "scout_output"
            ideas_json = scout_dir / "revive_scout_ideas.json"
            ideas_md = scout_dir / "new_topic_ideas.md"
            final_md = scout_dir / "Final_topic_suggest.md"

            cols = st.columns(3)
            for pth, col, label in [
                (ideas_json, cols[0], "revive_scout_ideas.json"),
                (ideas_md, cols[1], "new_topic_ideas.md"),
                (final_md, cols[2], "Final_topic_suggest.md"),
            ]:
                if pth.exists():
                    col.success(f"✔ {label}")
                    col.download_button(
                        label,
                        pth.read_bytes(),
                        file_name=pth.name,
                        use_container_width=True,
                    )
                else:
                    col.info(f"— {label} (없음)")

            # JSON 표로 확인
            if ideas_json.exists():
                try:
                    import pandas as _pd

                    data = json.loads(ideas_json.read_text(encoding="utf-8"))
                    if isinstance(data, list) and data:
                        df_ideas = _pd.DataFrame(
                            [
                                {
                                    "topic": it.get("topic", ""),
                                    "score": it.get("score", 0.0),
                                    "seed_pmids": ", ".join(
                                        (it.get("seed_pmids") or [])[:5]
                                    ),
                                    "seed_dois": ", ".join(
                                        (it.get("seed_dois") or [])[:5]
                                    ),
                                    "last_year_count": (
                                        lambda yc: yc.get(max(yc) if yc else None, None)
                                    )(it.get("year_counts", {})),
                                }
                                for it in data
                            ]
                        )
                        st.dataframe(df_ideas, use_container_width=True, height=320)
                except Exception as e:
                    st.warning(f"ideas.json 파싱 오류: {e}")

            # Final_topic_suggest 미리보기
            if final_md.exists():
                st.subheader("Final_topic_suggest.md (미리보기)")
                st.markdown(final_md.read_text(encoding="utf-8"))
