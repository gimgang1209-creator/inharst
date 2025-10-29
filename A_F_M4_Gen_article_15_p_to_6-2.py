# A_F_M4_Gen_article_15_p.py  (함수형 API + 기존 CLI 유지)
# -*- coding: utf-8 -*-

import os, re, json, argparse
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from openai import OpenAI, BadRequestError
except Exception:
    raise

# ---------------- CFG & Utils (기존 로직 대부분 유지) ----------------
CFG = {
    "model": "gpt-5",
    "style_rules": (
        "- 한국어 학술문체(객관적·정확한 서술)\n"
        "- IMRAD+결론+참고문헌 구조\n"
        "- 각 파트는 최소 단어수 충족. 부족하면 이어서 작성\n"
        "- 출처·수치·결과는 제공된 아티팩트를 우선 사용, 추측 금지\n"
        "- 표/그림은 '결과' 섹션에만 삽입. 다른 섹션에는 텍스트 설명만\n"
    ),
    "system_prompt": (
        "당신은 연구 자동화 시스템의 논문 작성 에이전트입니다. "
        "입력으로 제공된 M1(문헌), M2(분석), M3(도표) 아티팩트를 근거로 작성합니다. "
        "제공된 텍스트에 기반하고 미확인 정보는 꾸미지 마십시오."
    ),
    "section_instructions": {
        "introduction": "역할: 연구 배경, 필요성, 연구질문, 공헌. 최소 700단어. 표/그림 금지.",
        "methods": "역할: 데이터 출처, 변수, 집단, 전처리, 분석설계, 소프트웨어/버전. 최소 500단어. 표/그림 금지.",
        "results": "역할: 핵심 통계량, 효과크기, 신뢰구간, 유의성. 최소 500단어. 이 섹션에 한해 그림/표 삽입. 각 항목은 '캡션+2–4문장 설명'.",
        "discussion": "역할: 의미, 비교, 한계, 시사점. 향후과제, 최소 800단어. 표/그림 금지.",
        "conclusion": "역할: 최종 결과 기여. 최소 100단어. 표/그림 금지.",
        "references": "역할: 본문에서 실제 인용된 DOI만 [n] DOI 형식. 중복 금지.",
    },
}
SECTION_MIN = {
    "introduction": 700,
    "methods": 500,
    "results": 500,
    "discussion": 800,
    "conclusion": 100,
}

DOI_RE = re.compile(
    r"(?:doi:\s*|https?://doi\.org/)?(10\.\d{4,9}/[^\s\"\'\)\]]+)", re.I
)


def make_client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    if not key:
        raise RuntimeError("OPENAI_API_KEY가 비어 있습니다.")
    base = os.getenv("OPENAI_BASE", "").strip() or None
    org = os.getenv("OPENAI_ORG", "").strip() or None
    if base and org:
        return OpenAI(api_key=key, base_url=base, organization=org)
    if base:
        return OpenAI(api_key=key, base_url=base)
    if org:
        return OpenAI(api_key=key, organization=org)
    return OpenAI(api_key=key)


def _read(p: Path, n=8000) -> str:
    try:
        if not p.exists() or p.is_dir():
            return ""
        return p.read_text(encoding="utf-8", errors="ignore")[:n]
    except:
        return ""


def parse_section_min_words(s: str) -> dict:
    out = {}
    for pair in (s or "").split(","):
        if not pair.strip():
            continue
        k, v = pair.split("=")
        out[k.strip()] = int(v)
    return out


def _normalize_paragraphs(text: str) -> List[str]:
    t = re.sub(r"\r\n?", "\n", text or "")
    parts = re.split(r"\n\s*\n", t)
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        p = re.sub(r"\s*\n\s*", " ", p)
        out.append(p)
    return out


def _results_to_prose(text: str) -> str:
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    clean = []
    for ln in lines:
        if re.match(r"^\s*#{1,6}\s+", ln):
            continue
        if "|" in ln and re.search(r"\|\s*-{2,}", ln):
            continue
        if ln.strip().startswith("|"):
            continue
        ln = re.sub(r"^\s*[-*•]\s+", "", ln)
        ln = re.sub(r"^\s*\d+\.\s+", "", ln)
        ln = re.sub(r"^\s*\([a-zA-Z]\)\s+", "", ln)
        clean.append(ln)
    joined = "\n".join(clean)
    paras = _normalize_paragraphs(joined)
    merged = []
    for p in paras:
        if merged and len(p) < 60:
            merged[-1] = (merged[-1] + " " + p).strip()
        else:
            merged.append(p)
    return "\n\n".join(merged).strip()


def _extract_dois(*texts: str) -> List[str]:
    seen, out = set(), []
    big = "\n".join([t or "" for t in texts])
    for m in DOI_RE.finditer(big):
        doi = m.group(1).strip().rstrip(".,;)")
        if doi not in seen:
            seen.add(doi)
            out.append(doi)
    return out


def _replace_doi_citations(text: str, doi2num: Dict[str, int]) -> str:
    def repl(m):
        return f"[{doi2num.get(m.group(1).strip().rstrip('.,;)'),m.group(0))}]"

    t = DOI_RE.sub(repl, text or "")
    t = re.sub(r"\[(\d+)\]\s*\[(\d+)\]", r"[\1,\2]", t)
    for _ in range(3):
        t = re.sub(r"\[([\d,]+)\]\s*\[(\d+)\]", lambda m: f"[{m[1]},{m[2]}]", t)
    t = re.sub(r"\(\s*\[([\d, ]+)\]\s*\)", r"[\1]", t)
    t = re.sub(r"DOI\s*\[(\d+)\]", r"[\1]", t, flags=re.I)
    return t


# ---------------- OpenAI 호출 래퍼 (기존 로직 축약) ----------------
def _resp_text(resp) -> str:
    try:
        if getattr(resp, "output_text", None):
            return resp.output_text.strip()
        chunks = []
        for blk in getattr(resp, "output", []) or []:
            for c in getattr(blk, "content", []) or []:
                tv = getattr(getattr(c, "text", None), "value", None)
                if tv:
                    chunks.append(tv)
        if chunks:
            return "\n".join(chunks).strip()
    except:
        pass
    return (str(resp) or "").strip()


def _blocks_to_text(blocks: list[dict]) -> str:
    parts = []
    for b in blocks or []:
        if isinstance(b, dict) and b.get("type") == "input_text":
            parts.append(b.get("text", ""))
        else:
            # 혹시 dict가 아니거나 type이 없으면 문자열로 캐스팅
            parts.append(str(b))
    return "\n\n".join(parts).strip()


def _ask(
    client: OpenAI,
    model: str,
    system: str,
    user_blocks: list[dict],
    temp: float = 0.3,
    max_tokens: int = 1600,
) -> str:
    """
    1) responses API (신형) 시도
    2) 실패 시 chat.completions (구형/호환)로 폴백
    """
    # 1) responses API 시도 (SDK에 따라 temperature/max_output_tokens 직접 전달)
    try:
        r = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": system},
                {"role": "user", "content": user_blocks},
            ],
            temperature=temp,
            max_output_tokens=max_tokens,
        )
        return _resp_text(r)
    except Exception:
        pass

    # 2) 또 다른 responses 서명 시도 (최소 인자)
    try:
        r = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": system},
                {"role": "user", "content": user_blocks},
            ],
        )
        return _resp_text(r)
    except Exception:
        pass

    # 3) chat.completions로 폴백 (가장 호환성 높음)
    try:
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": _blocks_to_text(user_blocks)},
        ]
        r = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temp,
            max_tokens=max_tokens,
        )
        return (r.choices[0].message.content or "").strip()
    except Exception as e:
        return f"(LLM 호출 실패: {e})"


# ---------------- 핵심 함수형 API ----------------
def load_contexts(project: str, context_max_chars=8000) -> Dict[str, str]:
    proj = Path(project)
    A = proj / "agent_sub_result"
    R = proj / "analysis_result"
    intro = _read(A / "consolidated_report.md", context_max_chars)
    methods_hint = "데이터/모형 요약:\n" + _read(R / "hypothesis_validation.md", 8000)
    results_ctx = (
        _read(R / "final_analysis.md", 8000)
        + "\n"
        + _read(R / "final_analysis.json", 4000)
    )
    refs = _read(A / "filtered_articles_semantic.csv", 20000)
    return {
        "intro": intro,
        "methods": methods_hint,
        "results": results_ctx,
        "refs": refs,
    }


def draft_section(
    client, model, name_ko, key, ctx, min_words, max_tokens, max_cont=6, temp=0.3
) -> str:
    sysrules = CFG["system_prompt"] + "\n" + CFG["style_rules"]
    text = _ask(
        client,
        model,
        sysrules,
        [
            {"type": "input_text", "text": f"섹션: {name_ko}"},
            {
                "type": "input_text",
                "text": f"지시문:\n{CFG['section_instructions'][key]}",
            },
            {"type": "input_text", "text": f"참조 아티팩트 요약:\n{ctx[:10000]}"},
            {
                "type": "input_text",
                "text": f"제약: 표/그림 삽입 금지(결과 제외). 최소 {min_words} 단어.",
            },
        ],
        temp=temp,
        max_tokens=max_tokens,
    ).strip()

    # 이어붙이기(최소 단어수 보장)
    def wc(t):
        return len(re.findall(r"[가-힣]+|\w+", t or ""))

    guard = max_cont
    prev = -1
    while wc(text) < min_words and guard > 0 and len(text) != prev:
        prev = len(text)
        need = max(150, min_words - wc(text))
        addon = _ask(
            client,
            model,
            "같은 섹션을 자연스럽게 이어서 작성. 중복 최소화. 사실 기반 유지.",
            [
                {"type": "input_text", "text": f"배경 컨텍스트(요약):\n{ctx[:4000]}"},
                {"type": "input_text", "text": f"현재 초안(일부):\n{text[-6000:]}"},
                {
                    "type": "input_text",
                    "text": f"요구사항: 최소 {need}개 추가 단어. 목록/표/그림 금지(결과 제외). 한국어 학술문체.",
                },
            ],
            temp=temp,
            max_tokens=max_tokens,
        ).strip()
        if addon:
            text += "\n\n" + addon
        guard -= 1

    if key == "methods":
        text = "\n\n".join(_normalize_paragraphs(text))
    if key == "results":
        text = _results_to_prose(text)
    return text


def build_references(texts: List[str]) -> Tuple[Dict[str, int], str]:
    dois = _extract_dois(*texts)
    doi2num = {d: i + 1 for i, d in enumerate(dois)}
    refs = "\n".join(
        [f"[{n}] {d}" for d, n in sorted(doi2num.items(), key=lambda kv: kv[1])]
    )
    return doi2num, refs


def generate_sections(
    project: str,
    *,
    model: str | None = None,
    min_words: int = 500,
    max_tokens: int = 7000,
    temperature: float = 0.3,
) -> Dict:
    client = make_client()
    model = model or os.getenv("OPENAI_MODEL", "").strip() or CFG["model"]
    ctx = load_contexts(project)

    order = [
        ("서론", "introduction", ctx["intro"]),
        ("방법", "methods", ctx["methods"]),
        ("결과", "results", ctx["results"]),
        ("고찰", "discussion", ctx["intro"] + "\n\n" + ctx["results"]),
        ("결론", "conclusion", ctx["intro"] + "\n\n" + ctx["results"]),
    ]

    sections = {}
    for name_ko, key, context in order:
        target = min(max(SECTION_MIN.get(key, min_words), min_words), 100000)
        print(f"[생성] {name_ko}", flush=True)
        txt = draft_section(
            client, model, name_ko, key, context, target, max_tokens, temp=temperature
        )
        sections[key] = {"status": "confirmed", "text": txt}

    # DOI → [n] 치환 + 참고문헌
    full_before = "\n\n".join(
        [
            sections[k]["text"]
            for k in ["introduction", "methods", "results", "discussion", "conclusion"]
        ]
    )
    doi2num, refs = build_references([full_before, ctx["refs"]])
    for k in ["introduction", "methods", "results", "discussion", "conclusion"]:
        t = _replace_doi_citations(sections[k]["text"], doi2num)
        if k == "methods":
            t = "\n\n".join(_normalize_paragraphs(t))
        if k == "results":
            t = _results_to_prose(t)
        sections[k]["text"] = t
    sections["references"] = {"status": "confirmed", "text": refs}

    # 초록
    abs_text = _ask(
        client,
        model,
        "저널 스타일 한국어 논문 초록 편집기. 한 문단 200–300 단어. 수치/효과/신뢰구간 유지.",
        [
            {
                "type": "input_text",
                "text": "다음 본문 요약을 근거로 200–300 단어 한국어 초록(한 문단)을 작성하세요.",
            },
            {
                "type": "input_text",
                "text": (
                    sections["introduction"]["text"][:1000]
                    + "\n"
                    + sections["methods"]["text"][:900]
                    + "\n"
                    + sections["results"]["text"][:1200]
                    + "\n"
                    + sections["conclusion"]["text"][:600]
                ),
            },
        ],
        temp=0.2,
        max_tokens=min(800, max_tokens),
    ).strip()
    sections["abstract"] = {"status": "confirmed", "text": abs_text}

    return {
        "sections": sections,
        "options": {
            "min_words": min_words,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "model": model,
        },
    }


def set_korean_font(
    doc: Document, latin="Times New Roman", east_asia="Malgun Gothic", size=11
):
    style = doc.styles["Normal"]
    style.font.name = latin
    style.font.size = Pt(size)
    rpr = style._element.rPr
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), east_asia)
    rpr.append(rFonts)


def insert_md_table(doc: Document, md_text: str, caption: str, desc: str):
    # 간단 변환 (헤더 라인 없으면 열 추정)
    lines = [ln for ln in (md_text or "").splitlines() if ln.strip()]
    if not lines:
        return
    # 매우 러프한 파서
    cells = [
        [c.strip() for c in re.split(r"\s*\|\s*|\s{2,}|\t|,", ln.strip("|"))]
        for ln in lines
    ]
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, h in enumerate(cells[0]):
        hdr[j].text = h or f"Col{j+1}"
    for r in cells[1:]:
        row = table.add_row().cells
        for j, v in enumerate((r + [""] * ncols)[:ncols]):
            row[j].text = v
    cap = doc.add_paragraph(f"표: {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if desc:
        doc.add_paragraph(desc)
        doc.add_paragraph("")


def insert_figure(
    doc: Document,
    img_path: Path,
    caption: str,
    description: str,
    width_inches: float = 6.0,
):
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    cap_p = doc.add_paragraph(f"그림: {caption}")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if description:
        doc.add_paragraph(description)


def _caption_for(client, model, filename: str, context: str) -> Tuple[str, str]:
    out = _ask(
        client,
        model,
        "과학적 캡션·설명 작성 보조. 과장/추측 금지.",
        [
            {"type": "input_text", "text": f"파일명: {filename}"},
            {
                "type": "input_text",
                "text": "결과 섹션에 삽입할 항목의 '캡션'(한 문장)과 '설명'(2–4문장)을 작성.",
            },
            {"type": "input_text", "text": f"배경 컨텍스트(요약):\n{context[:4000]}"},
        ],
        temp=0.4,
        max_tokens=600,
    ).strip()
    m = re.search(r"(?:캡션|Caption)\s*[:：]\s*(.+)", out)
    cap = (
        m.group(1).strip()
        if m
        else (out.splitlines()[0].strip() if out else "결과 항목")
    )
    m2 = re.search(r"(?:설명|Explanation)\s*[:：]\s*(.+)", out, re.S)
    desc = m2.group(1).strip() if m2 else ""
    return cap, desc


def build_docx(project: str, state: Dict, output_path: str | None = None) -> Path:
    proj = Path(project)
    R = proj / "analysis_result"
    F = R / "figures"
    doc = Document()
    set_korean_font(doc)
    client = make_client()
    model = state.get("options", {}).get("model") or CFG["model"]

    # 제목(간단히 추출)
    intro = state["sections"].get("introduction", {}).get("text", "")
    results = state["sections"].get("results", {}).get("text", "")
    title = (
        _ask(
            client,
            model,
            "숙련된 학술 에디터. 과장 없이 정확·간결한 한 문장 한국어 제목만 출력.",
            [
                {"type": "input_text", "text": intro[:1200]},
                {"type": "input_text", "text": results[:1000]},
            ],
            temp=0.2,
            max_tokens=120,
        )
        .splitlines()[0]
        .strip(' "“”')
        or "연구 결과를 요약하는 한 문장 제목"
    )
    try:
        doc.core_properties.title = title
    except:
        pass

    # 제목/섹션
    def H(t):
        p = doc.add_heading(t, 1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    H("제목")
    doc.add_paragraph(title)

    order = [
        "abstract",
        "introduction",
        "methods",
        "results",
        "discussion",
        "conclusion",
        "references",
    ]
    ko = {
        "abstract": "초록",
        "introduction": "서론",
        "methods": "방법",
        "results": "결과",
        "discussion": "고찰",
        "conclusion": "결론",
        "references": "참고문헌",
    }

    img_paths = sorted(F.glob("*.png")) if F.exists() else []
    md_paths = sorted(F.glob("*.md")) if F.exists() else []
    md_texts = [_read(p, 100000) for p in md_paths]

    for key in order:
        sec = state["sections"].get(key, {"status": "canceled", "text": ""})
        if sec["status"] == "canceled":
            continue
        H(ko[key])
        if key == "results":
            # 문단 → 중간중간 그림/표 삽입
            paras = _normalize_paragraphs(sec["text"])
            cap_ctx = (sec["text"] + "\n\n" + _read(R / "final_analysis.md", 8000))[
                :8000
            ]
            # 간단한 균등 배치
            items = [("img", p) for p in img_paths] + [("md", t) for t in md_texts]
            if not items:
                for p in paras:
                    doc.add_paragraph(p)
            else:
                step = max(1, len(paras) // (len(items) + 1))
                idx = 0
                for i, p in enumerate(paras, start=1):
                    doc.add_paragraph(p)
                    if i % step == 0 and idx < len(items):
                        kind, payload = items[idx]
                        idx += 1
                        if kind == "img":
                            cap, desc = _caption_for(
                                client, model, payload.name, cap_ctx
                            )
                            insert_figure(doc, payload, cap, desc)
                        else:
                            cap, desc = _caption_for(client, model, "table.md", cap_ctx)
                            insert_md_table(doc, payload, cap, desc)
                # 남은 항목
                while idx < len(items):
                    kind, payload = items[idx]
                    idx += 1
                    if kind == "img":
                        cap, desc = _caption_for(client, model, payload.name, cap_ctx)
                        insert_figure(doc, payload, cap, desc)
                    else:
                        cap, desc = _caption_for(client, model, "table.md", cap_ctx)
                        insert_md_table(doc, payload, cap, desc)
        elif key == "references":
            for line in (sec["text"] or "").splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            for p in _normalize_paragraphs(sec["text"]):
                doc.add_paragraph(p)

    out_dir = proj / "m4_output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = (
        Path(output_path).resolve() if output_path else (out_dir / "manuscript.docx")
    )
    doc.save(str(out_path))

    # state 저장
    state_dump = {
        "sections": state["sections"],
        "options": state["options"],
        "title": title,
        "figures_inserted": [p.name for p in img_paths],
        "tables_inserted": [p.name for p in md_paths],
    }
    (out_dir / "m4_state.json").write_text(
        json.dumps(state_dump, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"[완료] {out_path}", flush=True)
    print(f"[제목] {title}", flush=True)
    return out_path


def generate_manuscript(
    project: str,
    *,
    model: str | None = None,
    min_words: int = 500,
    max_tokens: int = 7000,
    temperature: float = 0.3,
    output_path: str | None = None,
) -> Path:
    """외부(앱)에서 한 번에 호출하는 단일 API."""
    state = generate_sections(
        project,
        model=model,
        min_words=min_words,
        max_tokens=max_tokens,
        temperature=temperature,
    )
    return build_docx(project, state, output_path)


def run_interactive(args):
    client = make_client()
    model = args.model or os.getenv("OPENAI_MODEL", "").strip() or CFG["model"]
    ctx = load_contexts(args.project)

    order = [
        ("서론", "introduction", ctx["intro"]),
        ("방법", "methods", ctx["methods"]),
        ("결과", "results", ctx["results"]),
        ("고찰", "discussion", ctx["intro"] + "\n\n" + ctx["results"]),
        ("결론", "conclusion", ctx["intro"] + "\n\n" + ctx["results"]),
    ]

    state = {
        "sections": {},
        "options": {
            "min_words": args.min_words,
            "max_tokens": args.max_output_tokens,
            "temperature": args.temperature,
            "model": model,
        },
    }

    def emit(key: str, text: str):
        print(f"[[SECTION:start:{key}]]", flush=True)
        print(text, flush=True)
        print(f"[[SECTION:end:{key}]]", flush=True)

    # 섹션-by-섹션 생성 + 피드백 루프
    for name_ko, key, context in order:
        target = max(SECTION_MIN.get(key, args.min_words), args.min_words)
        print(f"[생성] {name_ko}", flush=True)

        # 최초 초안
        text = draft_section(
            client,
            model,
            name_ko,
            key,
            context,
            min_words=target,
            max_tokens=args.max_output_tokens,
            temp=args.temperature,
        )
        emit(key, text)

        # 피드백 루프
        while True:
            try:
                line = input().strip()
            except EOFError:
                line = ""
            if not line:
                # 엔터 → 확정 후 다음 섹션으로
                state["sections"][key] = {"status": "confirmed", "text": text}
                break
            if line == "취소":
                state["sections"][key] = {"status": "canceled", "text": text}
                break
            if line.startswith("수정:"):
                fb = line.split("수정:", 1)[1].strip()
                revised = _ask(
                    client,
                    model,
                    "섹션 개정: 피드백을 반영해 자연스럽게 고쳐쓰기. 중복 최소화, 사실 기반.",
                    [
                        {"type": "input_text", "text": f"[현재 초안]\n{text[-6000:]}"},
                        {"type": "input_text", "text": f"[피드백]\n{fb}"},
                        {
                            "type": "input_text",
                            "text": f"[제약] 최소 {target} 단어, 목록/표/그림 금지(결과 제외), 한국어 학술문체",
                        },
                    ],
                    temp=args.temperature,
                    max_tokens=min(3000, args.max_output_tokens),
                ).strip()
                if revised:
                    text = revised
                emit(key, text)
            else:
                print(
                    "[힌트] 엔터=확정/다음 · '수정: ...'=개정 · '취소'=건너뛰기",
                    flush=True,
                )

    # 참고문헌/초록/빌드
    full_before = "\n\n".join(
        [
            state["sections"][k]["text"]
            for k in ["introduction", "methods", "results", "discussion", "conclusion"]
            if k in state["sections"]
        ]
    )
    doi2num, refs = build_references([full_before, ctx["refs"]])
    for k in ["introduction", "methods", "results", "discussion", "conclusion"]:
        if k in state["sections"]:
            t = _replace_doi_citations(state["sections"][k]["text"], doi2num)
            if k == "methods":
                t = "\n\n".join(_normalize_paragraphs(t))
            if k == "results":
                t = _results_to_prose(t)
            state["sections"][k]["text"] = t
    state["sections"]["references"] = {"status": "confirmed", "text": refs}

    abs_text = _ask(
        client,
        model,
        "저널 스타일 한국어 논문 초록 편집기. 한 문단 200–300 단어. 수치/효과/신뢰구간 유지.",
        [
            {
                "type": "input_text",
                "text": "아래 본문 요약을 근거로 200–300 단어 초록(한 문단)을 작성.",
            },
            {
                "type": "input_text",
                "text": (
                    state["sections"].get("introduction", {}).get("text", "")[:1000]
                    + "\n"
                    + state["sections"].get("methods", {}).get("text", "")[:900]
                    + "\n"
                    + state["sections"].get("results", {}).get("text", "")[:1200]
                    + "\n"
                    + state["sections"].get("conclusion", {}).get("text", "")[:600]
                ),
            },
        ],
        temp=0.2,
        max_tokens=800,
    ).strip()
    state["sections"]["abstract"] = {"status": "confirmed", "text": abs_text}

    out_path = build_docx(args.project, state, args.output)
    print(f"[DONE]{out_path}", flush=True)


# ---------------- CLI (기존 호환) ----------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--project", type=str, default=".")
    ap.add_argument("--model", type=str, default=None)
    ap.add_argument("--min-words", type=int, default=500)
    ap.add_argument("--max-output-tokens", type=int, default=7000)
    ap.add_argument("--temperature", type=float, default=0.3)
    ap.add_argument("--output", type=str, default=None)
    ap.add_argument("--section-min-words", type=str, default="")

    args = ap.parse_args()

    if args.section_min_words:
        SECTION_MIN.update(parse_section_min_words(args.section_min_words))

    _ = generate_manuscript(
        args.project,
        model=args.model,
        min_words=args.min_words,
        max_tokens=args.max_output_tokens,
        temperature=args.temperature,
        output_path=args.output,
    )


if __name__ == "__main__":
    main()
