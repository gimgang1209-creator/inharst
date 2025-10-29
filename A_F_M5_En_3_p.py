#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import json
import argparse
import streamlit as st

# ---- common retry utilities (inserted by patch) ----
import time, requests

def http_get_with_retry(url, *, headers=None, params=None, data=None, timeout=15, tries=3, backoff=2):
    last_exc = None
    for i in range(tries):
        try:
            return http_get_with_retry(url, headers=headers, params=params, data=data, timeout=timeout)
        except Exception as e:
            last_exc = e
            if i < tries-1:
                time.sleep(backoff**i)
    raise last_exc

def between(val, lo, hi):
    try:
        return lo <= float(val) <= hi
    except Exception:
        return False
# ----------------------------------------------------

from pathlib import Path
from typing import Dict, List, Tuple

try:
    from openai import OpenAI, BadRequestError
except Exception:
    print("pip install openai")
    raise

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except Exception:
    print("pip install python-docx")
    raise

# ---------- DOI helpers ----------
DOI_EXTRACTOR = re.compile(r"10\.\d{4,9}/[^\s\"\'\)\]\}\.,;:]+", re.I)


def extract_first_doi_token(s: str) -> str:
    """문자열에서 첫 DOI만 안전 추출(소문자화 + 말단 구두점 제거)."""
    if not s:
        return ""
    m = DOI_EXTRACTOR.search(s)
    if not m:
        return ""
    return (m.group(0).lower().rstrip(").],;:")).strip()


def apa_citation_from_doi(doi: str, locale="en-US") -> str:
    """
    사용자가 요구한 간단/안정 버전.
    https://doi.org/…에 Accept 헤더로 APA 문자열을 요청.
    """
    import requests

    doi = (
        (doi or "")
        .strip()
        .lower()
        .replace("https://doi.org/", "")
        .replace("http://doi.org/", "")
    )
    url = f"https://doi.org/{doi}"
    headers = {
        "Accept": f"text/x-bibliography; style=apa; locale={locale}",
        "User-Agent": "apa-cite-script/1.0 (mailto:you@example.com)",
    }
    r = http_get_with_retry(url, headers=headers, timeout=15)
    r.raise_for_status()
    return r.text.strip()


def extract_dois_from_sections(
    sections: Dict[str, Dict[str, str]], extra: str = ""
) -> List[str]:
    """본문 섹션들(+옵션 문자열)에서 DOI를 등장 순서대로 추출(중복 제거)."""
    buf = []
    for k in [
        "abstract",
        "introduction",
        "methods",
        "results",
        "discussion",
        "conclusion",
        "references",
    ]:
        buf.append(sections.get(k, {}).get("text", ""))
    buf.append(extra or "")
    seen, out = set(), []
    for m in DOI_EXTRACTOR.finditer("\n".join(buf)):
        d = extract_first_doi_token(m.group(0))
        if d and d not in seen:
            seen.add(d)
            out.append(d)
    return out


def extract_dois_from_references_text(ref_text: str) -> List[str]:
    """
    이미 번호가 붙어있는 references.text에서 DOI만 깨끗이 추출.
    예) '[19] 10.1038/s41598-024-79055-1의', '[20] 10.1038/ejcn.2014.117,European' 등 처리.
    """
    seen, out = set(), []
    for line in (ref_text or "").splitlines():
        d = extract_first_doi_token(line)
        if d and d not in seen:
            seen.add(d)
            out.append(d)
    return out


# ---------- LLM helpers ----------
def get_resp_text(resp) -> str:
    try:
        if getattr(resp, "output_text", None):
            return resp.output_text
        out = []
        for blk in getattr(resp, "output", []) or []:
            for c in getattr(blk, "content", []) or []:
                tv = getattr(getattr(c, "text", None), "value", None)
                if tv:
                    out.append(tv)
        if out:
            return "\n".join(out).strip()
        out2 = getattr(resp, "output", None)
        if isinstance(out2, list) and out2 and getattr(out2[0], "content", None):
            cand = out2[0].content[0]
            tv = getattr(getattr(cand, "text", None), "value", None)
            if tv:
                return tv.strip()
    except Exception:
        pass
    s = str(resp)
    if "object='response'" in s and "output_tokens" in s:
        return ""
    return s.strip()


def safe_text(resp) -> str:
    t = (get_resp_text(resp) or "").strip()
    if t.startswith("Response(") and "object='response'" in t:
        return ""
    return t


SYS_TRANSLATE = (
    "You are a professional academic translator. Translate Korean scientific prose to formal, idiomatic academic English. "
    "Preserve meaning, numbers, units, effect sizes, confidence intervals, and bracketed citation indices like [3] as-is. "
    "Do not invent references or numbers. No lists unless present; prefer cohesive paragraphs."
)


def call_response(
    client: OpenAI,
    model: str,
    input_payload,
    temperature: float,
    max_output_tokens: int,
):
    try:
        return client.responses.create(
            model=model,
            input=input_payload,
            inference_config={
                "temperature": temperature,
                "max_output_tokens": max_output_tokens,
            },
        )
    except TypeError:
        pass
    try:
        return client.responses.create(
            model=model, input=input_payload, max_output_tokens=max_output_tokens
        )
    except BadRequestError:
        pass
    except TypeError:
        pass
    try:
        return client.responses.create(
            model=model,
            input=input_payload,
            temperature=temperature,
            max_output_tokens=max_output_tokens,
        )
    except BadRequestError as e3:
        em = getattr(e3, "message", "") or ""
        if "Unsupported parameter" in em and "temperature" in em:
            return client.responses.create(
                model=model, input=input_payload, max_output_tokens=max_output_tokens
            )
        raise
    except TypeError:
        return client.responses.create(model=model, input=input_payload)


def translate_block(
    client, model, text: str, temperature=0.1, max_output_tokens=2000
) -> str:
    if not (text or "").strip():
        return ""
    resp = call_response(
        client,
        model,
        input_payload=[
            {"role": "system", "content": SYS_TRANSLATE},
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": "Translate to academic English:"},
                    {"type": "input_text", "text": (text or "")[:12000]},
                ],
            },
        ],
        temperature=temperature,
        max_output_tokens=max_output_tokens,
    )
    t = safe_text(resp).strip()
    return t or text


def refine_abstract_english(
    client, model, abstract_en: str, temperature=0.1, max_output_tokens=800
) -> str:
    if not (abstract_en or "").strip():
        return ""
    resp = call_response(
        client,
        model,
        input_payload=[
            {
                "role": "system",
                "content": "Journal abstract polisher. Output a single 200–300-word paragraph in formal academic English. No headings, bullets, tables, or figure mentions.",
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": "Polish this into one journal-style paragraph (200–300 words):",
                    },
                    {"type": "input_text", "text": abstract_en[:4000]},
                ],
            },
        ],
        temperature=temperature,
        max_output_tokens=max_output_tokens,
    )
    t = safe_text(resp).strip()
    return t or abstract_en


def prompt_caption_en(
    client,
    model,
    file_name: str,
    context_hint: str,
    temperature=0.2,
    max_output_tokens=400,
) -> Tuple[str, str]:
    resp = call_response(
        client,
        model,
        input_payload=[
            {
                "role": "system",
                "content": "You write concise scientific figure/table captions and 2–4 sentence descriptions in English. No overclaims.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": f"File: {file_name}"},
                    {
                        "type": "input_text",
                        "text": "Return: Caption: <one sentence>. Description: <2–4 sentences>.",
                    },
                    {
                        "type": "input_text",
                        "text": f"Context (summary):\n{(context_hint or '')[:4000]}",
                    },
                ],
            },
        ],
        temperature=temperature,
        max_output_tokens=max_output_tokens,
    )
    out = safe_text(resp).strip()
    cap = None
    desc = None
    m = re.search(r"(?:Caption)\s*[:：]\s*(.+)", out)
    if m:
        cap = m.group(1).strip()
    m2 = re.search(r"(?:Description)\s*[:：]\s*(.+)", out, re.S)
    if m2:
        desc = m2.group(1).strip()
    if not cap:
        parts = re.split(r"[\n\.]\s+", out, maxsplit=1)
        cap = parts[0].strip() if parts else "Result item"
        desc = parts[1].strip() if len(parts) > 1 else ""
    return cap, desc


# ---------- DOCX helpers ----------
def set_font_en(
    document: Document, latin="Times New Roman", east_asia="Malgun Gothic", size=11
):
    style = document.styles["Normal"]
    style.font.name = latin
    style.font.size = Pt(size)
    rpr = style._element.rPr
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), east_asia)
    rpr.append(rFonts)


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def normalize_paragraphs(text: str) -> List[str]:
    t = re.sub(r"\r\n?", "\n", text or "")
    paras = re.split(r"\n\s*\n", t)
    out = []
    for p in paras:
        p = p.strip()
        if not p:
            continue
        p = re.sub(r"\s*\n\s*", " ", p)
        out.append(p.strip())
    return out


def insert_md_table(doc: Document, md_text: str, caption: str, desc: str):
    lines = [ln.rstrip() for ln in (md_text or "").splitlines() if ln.strip()]
    header, rows, i = [], [], 0
    while i < len(lines):
        ln = lines[i]
        if "|" in ln:
            cells = [c.strip() for c in ln.strip("|").split("|")]
            if i + 1 < len(lines) and re.match(r"^\s*\|?:?-{2,}", lines[i + 1]):
                header = cells
                i += 2
                break
        i += 1
    while i < len(lines):
        ln = lines[i]
        if "|" in ln:
            cells = [c.strip() for c in ln.strip("|").split("|")]
            if not header or len(cells) == len(header):
                rows.append(cells)
        i += 1
    if not header and not rows:
        return
    ncols = len(header) if header else len(rows[0])
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, h in enumerate(header or [f"Col{j+1}" for j in range(ncols)]):
        hdr[j].text = h
    for r in rows:
        row = table.add_row().cells
        for j, v in enumerate(r[:ncols]):
            row[j].text = v
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if desc:
        doc.add_paragraph(desc)
        doc.add_paragraph("")


def insert_figure(
    doc: Document, img_path: Path, caption: str, description: str, width_inches=6.0
):
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if description:
        doc.add_paragraph(description)


def distribute_positions(n_paras: int, n_items: int) -> List[int]:
    if n_paras <= 0 or n_items <= 0:
        return []
    step = (n_paras + 1) / (n_items + 1)
    return [max(0, min(n_paras, int(round(step * (i + 1))))) for i in range(n_items)]


def interleave_results_en(
    doc: Document,
    result_paragraphs: List[str],
    img_paths: List[Path],
    md_texts: List[str],
    caption_context: str,
    client: OpenAI,
    model: str,
    temperature: float,
    max_output_tokens: int,
):
    items = [("img", p) for p in img_paths] + [("md", t) for t in md_texts]
    positions = distribute_positions(len(result_paragraphs), len(items))
    fig_idx, tab_idx, cur = 1, 1, 0
    for i, para in enumerate(result_paragraphs, start=1):
        doc.add_paragraph(para)
        while cur < len(positions) and positions[cur] == i:
            kind, payload = items[cur]
            if kind == "img":
                p = payload
                cap, desc = prompt_caption_en(
                    client,
                    model,
                    p.name,
                    caption_context,
                    temperature=min(0.3, temperature + 0.1),
                    max_output_tokens=max_output_tokens // 4,
                )
                insert_figure(doc, p, f"Figure {fig_idx}. {cap}", desc)
                fig_idx += 1
            else:
                md = payload
                cap, desc = prompt_caption_en(
                    client,
                    model,
                    "table.md",
                    caption_context,
                    temperature=min(0.3, temperature + 0.1),
                    max_output_tokens=max_output_tokens // 4,
                )
                insert_md_table(doc, md, f"Table {tab_idx}. {cap}", desc)
                tab_idx += 1
            cur += 1
    while cur < len(items):
        kind, payload = items[cur]
        if kind == "img":
            p = payload
            cap, desc = prompt_caption_en(
                client,
                model,
                p.name,
                caption_context,
                temperature=min(0.3, temperature + 0.1),
                max_output_tokens=max_output_tokens // 4,
            )
            insert_figure(doc, p, f"Figure {fig_idx}. {cap}", desc)
            fig_idx += 1
        else:
            md = payload
            cap, desc = prompt_caption_en(
                client,
                model,
                "table.md",
                caption_context,
                temperature=min(0.3, temperature + 0.1),
                max_output_tokens=max_output_tokens // 4,
            )
            insert_md_table(doc, md, f"Table {tab_idx}. {cap}", desc)
            tab_idx += 1
        cur += 1


# ---------- Title helpers ----------
KOREAN_RE = re.compile(r"[\u3131-\u318E\uAC00-\uD7A3]")  # 한글 자모/가-힣


def is_korean(s: str) -> bool:
    return bool(s and KOREAN_RE.search(s))


def translate_title_ko_to_en(client, model: str, text: str) -> str:
    """제목 전용 번역: 한 줄, 과도한 의역 금지, 고유명사/약어 보존"""
    if not text.strip():
        return ""
    resp = call_response(
        client,
        model,
        input_payload=[
            {
                "role": "system",
                "content": (
                    "You are a professional academic translator. "
                    "Translate the Korean paper title into concise, formal academic English. "
                    "Return ONE line only, no quotes, no extra text. "
                    "Preserve proper nouns and acronyms."
                ),
            },
            {"role": "user", "content": [{"type": "input_text", "text": text[:300]}]},
        ],
        temperature=0.1,
        max_output_tokens=80,
    )
    t = safe_text(resp).strip()
    return t.strip(' "“”')


def heuristic_title_from_sections_en(sections_en: Dict[str, Dict[str, str]]) -> str:
    intro = sections_en.get("introduction", {}).get("text", "")
    results = sections_en.get("results", {}).get("text", "")
    base = intro[:1500] + "\n" + results[:1500]
    tokens = re.findall(r"[A-Za-z][A-Za-z\-]{2,}", base)
    key = (
        " ".join(tokens[:12]).strip()
        or "A concise one-line title summarizing the study"
    )
    return (key[:140]).strip(" .-")


# ---------- Review helpers ----------
def review_loop(name: str, text: str, non_interactive=False) -> Tuple[str, str]:
    """
    Streamlit 환경 전용 피드백 루프.
    non_interactive=True면 자동 확정.
    """
    if non_interactive:
        return "confirmed", text

    st.markdown(f"### 🧾 {name} Preview")
    st.text_area(f"{name} (영문 초안)", text, height=280, key=f"draft_{name}")
    instruction = st.text_area(f"{name} 수정 지시문 입력 (예: Make it more concise...)", key=f"instr_{name}")

    col1, col2, col3 = st.columns(3)
    action = None
    with col1:
        if st.button(f"{name} 확정", key=f"confirm_{name}"):
            action = "confirmed"
    with col2:
        if st.button(f"{name} 수정 반영", key=f"revise_{name}"):
            action = "revise"
    with col3:
        if st.button(f"{name} 취소", key=f"cancel_{name}"):
            action = "canceled"

    if action == "revise" and instruction.strip():
        return "revise", instruction.strip()
    elif action == "canceled":
        return "canceled", ""
    elif action == "confirmed":
        return "confirmed", text
    else:
        st.info("버튼을 눌러 진행하세요.")
        st.stop()


def apply_revision_en(
    client, model, draft_text: str, instruction: str, max_output_tokens: int
) -> str:
    resp = call_response(
        client,
        model,
        input_payload=[
            {
                "role": "system",
                "content": "Apply the user's editing instruction conservatively to the English academic draft. Keep facts and logic.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": "Draft:\n" + draft_text},
                    {"type": "input_text", "text": "Instruction:\n" + instruction},
                ],
            },
        ],
        temperature=0.2,
        max_output_tokens=max_output_tokens // 2,
    )
    t = safe_text(resp).strip()
    return t or draft_text


# ---------- Main ----------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--project",
        type=str,
        default=".",
        help="Project root containing m4_output/m4_state.json",
    )
    ap.add_argument("--model", type=str, default="gpt-4o-mini")
    ap.add_argument("--locale", type=str, default="en-US")
    ap.add_argument(
        "--output",
        type=str,
        default=None,
        help="Output .docx (default: m4_output/manuscript_en.docx)",
    )
    ap.add_argument(
        "--non-interactive", action="store_true", help="Skip interactive edits"
    )
    ap.add_argument("--max-output-tokens", type=int, default=6000)
    ap.add_argument("--temperature", type=float, default=0.2)
    args = ap.parse_args()

    proj = Path(args.project).resolve()
    out_dir = proj / "m4_output"
    out_dir.mkdir(parents=True, exist_ok=True)
    state_path = out_dir / "m4_state.json"
    if not state_path.exists():
        raise FileNotFoundError(f"State file not found: {state_path}")

    # 항상 선행 정의: figures 경로/리스트 (NameError 방지)
    R = proj / "analysis_result"
    FDIR = R / "figures"

    # FDIR.mkdir(parents=True, exist_ok=True)  # 필요 시 활성화
    def read_text(p: Path, enc="utf-8"):
        try:
            return p.read_text(encoding=enc)
        except Exception:
            return ""

    img_paths = sorted(FDIR.glob("*.png")) if FDIR.exists() else []
    md_paths = sorted(FDIR.glob("*.md")) if FDIR.exists() else []
    md_texts = [read_text(p) for p in md_paths]

    state = json.loads(state_path.read_text(encoding="utf-8"))
    sections = state.get("sections", {})
    title_ko = (state.get("title", "") or "").strip()  # 사용자가 명시: 항상 여기에 있음

    # OpenAI client (환경변수 권장)
    client = OpenAI()

    # 1) Translate sections to English
    en_sections = {}
    order = [
        "abstract",
        "introduction",
        "methods",
        "results",
        "discussion",
        "conclusion",
    ]

    for key in order:
        txt = sections.get(key, {}).get("text", "")
        en = translate_block(
            client, args.model, txt,
            temperature=args.temperature, max_output_tokens=args.max_output_tokens
        )
        if key == "abstract":
            en = refine_abstract_english(
                client, args.model, en,
                temperature=0.1, max_output_tokens=min(800, args.max_output_tokens)
            )

        # 🔸 이전 개정본이 있으면 이어서 표시
        ss_key = f"sec_{key}"
        if ss_key in st.session_state and st.session_state[ss_key]:
            en = st.session_state[ss_key]

        # 🔁 확정될 때까지 반복
        while True:
            status, payload = review_loop(key.capitalize(), en, args.non_interactive)

            if status == "revise":
                # 개정 적용
                en = apply_revision_en(client, args.model, en, payload, args.max_output_tokens)
                # 🔸 개정본을 세션에 저장하고 즉시 리런(미리보기 갱신)
                st.session_state[ss_key] = en
                st.rerun()

            elif status == "canceled":
                en_sections[key] = {"status": "canceled", "text": ""}
                break

            else:  # confirmed
                en_sections[key] = {"status": "confirmed", "text": en}
                # 🔸 확정되면 세션 값 정리(선택)
                st.session_state.pop(ss_key, None)
                break














    # 2) Title in English (간단 소스: state["title"]만 사용)
    title_en = ""
    if title_ko:
        if is_korean(title_ko):
            title_en = translate_title_ko_to_en(client, args.model, title_ko)
            if is_korean(title_en):  # 안전 재시도
                title_en = (
                    translate_block(
                        client,
                        args.model,
                        title_ko,
                        temperature=0.1,
                        max_output_tokens=200,
                    )
                    .splitlines()[0]
                    .strip(' "“”')
                )
        else:
            title_en = title_ko.splitlines()[0].strip(' "“”')
    if not title_en:
        title_en = heuristic_title_from_sections_en(en_sections)
    
    while True:
        status, payload = review_loop("Title", title_en, args.non_interactive)
        if status == "revise":
            title_en = (
                apply_revision_en(client, args.model, title_en, payload, 600)
                .splitlines()[0]
                .strip(' "“”')
            )
            continue  # 개정 후 다시 미리보기로 돌아감
        elif status == "canceled":
            break
        else:  # confirmed
            break

    title_en = (title_en or "A concise one-line title summarizing the study").strip()

    # 3) APA references from DOIs
    A = proj / "agent_sub_result"
    ctx_refs_raw = read_text(A / "filtered_articles_semantic.csv")

    # 우선 JSON의 원본 references.text에서 DOI만 추출
    existing_refs_text = (sections.get("references", {}) or {}).get("text", "")
    doi_list = extract_dois_from_references_text(existing_refs_text)

    # 비어있다면 섹션 전체에서 스캔
    if not doi_list:
        doi_list = extract_dois_from_sections(en_sections, extra=ctx_refs_raw)

    citations, apa_map = [], {}
    for i, d in enumerate(doi_list, start=1):
        try:
            apa = apa_citation_from_doi(d, locale=args.locale)
        except Exception:
            apa = f"Unresolved DOI: https://doi.org/{d}"
        apa_map[d] = apa
        citations.append(f"[{i}] {apa}")  # 요구사항: [n] 프리픽스

    refs_text = "\n".join(citations)
    en_sections["references"] = {"status": "confirmed", "text": refs_text}

    # 4) Build English Word doc
    doc = Document()
    set_font_en(doc)

    # Title FIRST
    try:
        doc.core_properties.title = title_en
    except Exception:
        pass
    add_heading(doc, "Title", level=1)
    doc.add_paragraph(title_en)

    headings_en = {
        "abstract": "Abstract",
        "introduction": "Introduction",
        "methods": "Methods",
        "results": "Results",
        "discussion": "Discussion",
        "conclusion": "Conclusion",
        "references": "References",
    }
    final_order = [
        "abstract",
        "introduction",
        "methods",
        "results",
        "discussion",
        "conclusion",
        "references",
    ]

    for key in final_order:
        sec = en_sections.get(key, {"status": "canceled", "text": ""})
        if sec["status"] == "canceled":
            continue
        add_heading(doc, headings_en[key], level=1)
        if key == "results":
            # interleave figures/tables again, captions in EN
            caption_ctx = (sec["text"] + "\n\n" + read_text(R / "final_analysis.md"))[
                :8000
            ]
            paras = normalize_paragraphs(sec["text"])
            interleave_results_en(
                doc,
                paras,
                img_paths,
                md_texts,
                caption_ctx,
                client,
                args.model,
                args.temperature,
                args.max_output_tokens,
            )
        elif key == "references":
            # 각 항목이 이미 [n] APA 문자열이므로 한 줄씩 추가
            for line in (sec["text"] or "").splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        else:
            for p in normalize_paragraphs(sec["text"]):
                doc.add_paragraph(p)

    out_path = (
        Path(args.output).resolve() if args.output else out_dir / "manuscript_en.docx"
    )
    doc.save(str(out_path))

    # 5) Save new state (NameError 방지: 앞에서 항상 정의됨)
    new_state = {
        "title_en": title_en,
        "sections_en": en_sections,
        "apa_locale": args.locale,
        "dois": doi_list,
        "apa_map": apa_map,
        "figures_inserted": [p.name for p in img_paths],
        "tables_inserted": [p.name for p in md_paths],
        "source_state_file": str(state_path),
    }
    (out_dir / "m4_state_en.json").write_text(
        json.dumps(new_state, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(f"[DONE] {out_path}")
    print(f"[TITLE] {title_en}")
    print(f"[REFERENCES] {len(doi_list)} items (APA, {args.locale})")


if __name__ == "__main__":
    main()
